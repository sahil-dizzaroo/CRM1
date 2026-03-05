from fastapi import APIRouter, Depends, HTTPException, WebSocket, WebSocketDisconnect, Header, Query, UploadFile, File, Form, Request
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.security import OAuth2PasswordRequestForm
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, or_, and_, delete, update, func
from sqlalchemy.orm import selectinload
from typing import Optional, List
from uuid import UUID
from datetime import timedelta, datetime, timezone
import asyncio
import os
import hashlib
import shutil
from pathlib import Path
import secrets
from app.db import get_db, init_db
from app import crud, schemas
from app.models import MessageDirection, MessageStatus, MessageChannel, ConversationAccessLevel, ThreadAttachment, Attachment, Conversation, ChatDocument, PrimarySiteStatus, UserRoleAssignment, Site, Study, StudySite, SiteStatus, SiteWorkflowStep, SiteDocument, WorkflowStepName, StepStatus, DocumentCategory, DocumentType, ReviewStatus, ProjectFeasibilityCustomQuestion, FeasibilityRequest, FeasibilityResponse, FeasibilityRequestStatus, FeasibilityAttachment, Agreement, AgreementComment, AgreementStatus, CommentType, StudyTemplate, AgreementDocument, TemplateType, AgreementSignedDocument, SiteProfile
from app.websocket_manager import manager
from app.config import settings
from app.auth import create_access_token, get_password_hash, get_current_user, get_current_user_optional, ACCESS_TOKEN_EXPIRE_MINUTES
from app.site_status_service import (
    get_study_status_summary,
    get_country_site_counts,
    get_sites_by_status,
    get_site_status_detail,
)
from app.api.v1.endpoints.clinical_workflow import (
    get_or_create_study_site,
    generate_cda_document_html,
    get_cda_document_path,
)
from app.ai_service import ai_service
from app.services.agreement_service import (
    change_agreement_status as change_agreement_status_service,
    get_next_allowed_status as get_next_allowed_status_service,
    check_can_upload_new_version,
    filter_versions_by_role,
    is_user_internal,
    can_create_comment_type,
    filter_comments_by_role,
    get_editor_permissions,
    get_agreement_permissions,
    create_agreement_notice,
    AgreementStatus as AgreementStatusEnum,
)
import uuid
import logging
# Tasks imported where needed to avoid circular imports


router = APIRouter(tags=["Legal"])
logger = logging.getLogger(__name__)

@router.post("/webhooks/mock_provider")
async def webhook_mock_provider(
    payload: schemas.WebhookPayload,
    x_mock_token: Optional[str] = Header(None, alias="X-MOCK-TOKEN")
):
    # Validate token
    if x_mock_token != settings.mock_provider_token:
        raise HTTPException(status_code=401, detail="Invalid token")
    
    # Enqueue webhook processing task
    from app.workers.tasks import process_webhook_task
    process_webhook_task.delay(payload.dict())
    
    return {"status": "accepted"}


@router.get("/sites/{site_id}/documents", response_model=List[schemas.SiteDocumentResponse])
async def get_site_documents(
    site_id: str,
    category: Optional[str] = Query(None),
    document_type: Optional[str] = Query(None, description="Filter by document type: 'sponsor' or 'site'"),
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    List all documents for a site.
    Optional category and document_type filters.
    """
    # Resolve site
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()
    
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")
    
    # Query documents
    query = select(SiteDocument).where(SiteDocument.site_id == site.id)
    if category:
        try:
            category_enum = DocumentCategory(category)
            query = query.where(SiteDocument.category == category_enum)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid category: {category}")
    
    if document_type:
        try:
            doc_type_enum = DocumentType(document_type)
            query = query.where(SiteDocument.document_type == doc_type_enum)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid document_type: {document_type}")
    
    query = query.order_by(SiteDocument.uploaded_at.desc())
    result = await db.execute(query)
    documents = result.scalars().all()
    
    return [
        {
            "id": doc.id,
            "site_id": str(doc.site_id),
            "category": doc.category.value,
            "file_name": doc.file_name,
            "content_type": doc.content_type,
            "size": doc.size,
            "uploaded_by": doc.uploaded_by,
            "uploaded_at": doc.uploaded_at,
            "description": doc.description,
            "metadata": doc.document_metadata or {},
            "document_type": doc.document_type.value if doc.document_type else None,
            "review_status": doc.review_status.value if doc.review_status else None,
            "tmf_filed": doc.tmf_filed or "false",
        }
        for doc in documents
    ]


@router.post("/sites/{site_id}/documents", response_model=schemas.SiteDocumentResponse)
async def upload_site_document(
    site_id: str,
    file: UploadFile = File(...),
    category: str = Form(...),
    description: Optional[str] = Form(None),
    document_type: Optional[str] = Form(None, description="Document type: 'sponsor' or 'site'. Defaults to 'site'."),
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Upload a document for a site.
    Document persists regardless of site status changes.
    """
    # Resolve site
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()
    
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")
    
    # Validate category
    try:
        category_enum = DocumentCategory(category)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid category: {category}")
    
    # Validate and set document_type
    doc_type_enum = DocumentType.SITE  # Default to SITE
    if document_type:
        try:
            doc_type_enum = DocumentType(document_type)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid document_type: {document_type}")
    
    # Set review_status: sponsor documents don't need review, site documents start as pending
    review_status_enum = None
    if doc_type_enum == DocumentType.SITE:
        review_status_enum = ReviewStatus.PENDING
    
    # Create upload directory if it doesn't exist
    upload_dir = Path(settings.upload_dir) / "site_documents"
    upload_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate unique filename
    file_ext = Path(file.filename).suffix if file.filename else ""
    file_id = uuid.uuid4()
    file_name = f"{file_id}{file_ext}"
    file_path = upload_dir / file_name
    
    # Save file
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        file_size = file_path.stat().st_size
        
        # Create document record
        document = SiteDocument(
            site_id=site.id,
            category=category_enum,
            file_path=str(file_path),
            file_name=file.filename or "unknown",
            content_type=file.content_type or "application/octet-stream",
            size=file_size,
            uploaded_by=current_user.get("user_id") if current_user else None,
            description=description,
            document_type=doc_type_enum,
            review_status=review_status_enum,
            tmf_filed='false',
        )
        db.add(document)
        await db.commit()
        await db.refresh(document)
        
        return {
            "id": document.id,
            "site_id": str(document.site_id),
            "category": document.category.value,
            "file_name": document.file_name,
            "content_type": document.content_type,
            "size": document.size,
            "uploaded_by": document.uploaded_by,
            "uploaded_at": document.uploaded_at,
            "description": document.description,
            "metadata": document.document_metadata or {},
            "document_type": document.document_type.value if document.document_type else None,
            "review_status": document.review_status.value if document.review_status else None,
            "tmf_filed": document.tmf_filed or "false",
        }
        
    except Exception as e:
        # Clean up file if database operation fails
        if file_path.exists():
            file_path.unlink()
        raise HTTPException(status_code=500, detail=f"Failed to upload document: {str(e)}")


@router.get("/sites/{site_id}/documents/{document_id}/download")
async def download_site_document(
    site_id: str,
    document_id: UUID,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Download a site document.
    """
    # Resolve site
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()
    
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")
    
    # Get document
    result = await db.execute(
        select(SiteDocument).where(
            SiteDocument.id == document_id,
            SiteDocument.site_id == site.id
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    file_path = Path(document.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on server")
    
    return FileResponse(
        path=str(file_path),
        filename=document.file_name,
        media_type=document.content_type,
    )


async def sendDocumentToTMF(document_id: UUID, db: AsyncSession) -> bool:
    """
    Placeholder function for TMF integration.
    In the future, this will integrate with the actual TMF system.
    
    Args:
        document_id: The UUID of the document to send to TMF
        db: Database session
        
    Returns:
        bool: True if successful, False otherwise
    """
    # TODO: Implement actual TMF integration
    # For now, this is a placeholder that simulates TMF filing
    print(f"[TMF PLACEHOLDER] Would send document {document_id} to TMF system")
    return True


@router.post("/sites/{site_id}/documents/{document_id}/approve", response_model=schemas.SiteDocumentResponse)
async def approve_site_document(
    site_id: str,
    document_id: UUID,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Approve a site-uploaded document and send it to TMF.
    Only works for documents with document_type='site' and review_status='pending'.
    """
    # Resolve site
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()
    
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")
    
    # Get document
    result = await db.execute(
        select(SiteDocument).where(
            SiteDocument.id == document_id,
            SiteDocument.site_id == site.id
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Validate that this is a site-uploaded document
    if not document.document_type or document.document_type != DocumentType.SITE:
        raise HTTPException(status_code=400, detail="Only site-uploaded documents can be approved")
    
    # Update review status
    document.review_status = ReviewStatus.APPROVED
    
    # Send to TMF (placeholder)
    tmf_success = await sendDocumentToTMF(document.id, db)
    if tmf_success:
        document.tmf_filed = 'true'
    else:
        # Even if TMF fails, mark as approved (TMF can be retried later)
        pass
    
    await db.commit()
    await db.refresh(document)
    
    return {
        "id": document.id,
        "site_id": str(document.site_id),
        "category": document.category.value,
        "file_name": document.file_name,
        "content_type": document.content_type,
        "size": document.size,
        "uploaded_by": document.uploaded_by,
        "uploaded_at": document.uploaded_at,
        "description": document.description,
        "metadata": document.document_metadata or {},
        "document_type": document.document_type.value if document.document_type else None,
        "review_status": document.review_status.value if document.review_status else None,
        "tmf_filed": document.tmf_filed or "false",
    }


@router.post("/sites/{site_id}/documents/{document_id}/reject", response_model=schemas.SiteDocumentResponse)
async def reject_site_document(
    site_id: str,
    document_id: UUID,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Reject a site-uploaded document.
    Only works for documents with document_type='site' and review_status='pending'.
    """
    # Resolve site
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()
    
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")
    
    # Get document
    result = await db.execute(
        select(SiteDocument).where(
            SiteDocument.id == document_id,
            SiteDocument.site_id == site.id
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Validate that this is a site-uploaded document
    if not document.document_type or document.document_type != DocumentType.SITE:
        raise HTTPException(status_code=400, detail="Only site-uploaded documents can be rejected")
    
    # Update review status
    document.review_status = ReviewStatus.REJECTED
    
    await db.commit()
    await db.refresh(document)
    
    return {
        "id": document.id,
        "site_id": str(document.site_id),
        "category": document.category.value,
        "file_name": document.file_name,
        "content_type": document.content_type,
        "size": document.size,
        "uploaded_by": document.uploaded_by,
        "uploaded_at": document.uploaded_at,
        "description": document.description,
        "metadata": document.document_metadata or {},
        "document_type": document.document_type.value if document.document_type else None,
        "review_status": document.review_status.value if document.review_status else None,
        "tmf_filed": document.tmf_filed or "false",
    }


# ---------------------------------------------------------------------------
# Feasibility Questionnaire Endpoints
# ---------------------------------------------------------------------------

@router.post("/sites/{site_id}/cda/sign-internal")
async def sign_cda_internal(
    site_id: str,
    study_id: Optional[str] = Query(
        None,
        description="Study identifier for study-specific CDA (can be UUID or external study id, e.g. ASLAN001-009)",
    ),
    signer_name: str = Form(...),
    signer_title: str = Form(...),
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Internal (Dizzaroo) signing of CDA.
    Updates the CDA Execution step_data with internal signer information.
    """
    # Resolve site
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()

    if not site:
        raise HTTPException(status_code=404, detail="Site not found")

    # Resolve study_site for study-specific CDA if study_id is provided
    study_site: Optional[StudySite] = None
    if study_id:
        # Try UUID first, then external identifiers; if not found, fall back to site-level CDA
        study_obj: Optional[Study] = None
        try:
            study_uuid = UUID(str(study_id))
            study_result = await db.execute(select(Study).where(Study.id == study_uuid))
            study_obj = study_result.scalar_one_or_none()
        except (ValueError, TypeError):
            # Fallback to external identifiers: study.study_id or study.name
            study_result = await db.execute(
                select(Study).where(
                    or_(
                        Study.study_id == str(study_id),
                        Study.name == str(study_id),
                    )
                )
            )
            study_obj = study_result.scalar_one_or_none()

        if study_obj:
            study_site = await get_or_create_study_site(db, study_obj.id, site.id)
        # If study cannot be resolved, we simply don't scope by study_site and use site-level CDA

    # Get CDA Execution step (study-specific if study_site exists, otherwise site-level)
    if study_site:
        step_result = await db.execute(
            select(SiteWorkflowStep).where(
                SiteWorkflowStep.study_site_id == study_site.id,
                SiteWorkflowStep.step_name == WorkflowStepName.CDA_EXECUTION.value,
            )
        )
    else:
        step_result = await db.execute(
            select(SiteWorkflowStep).where(
                SiteWorkflowStep.site_id == site.id,
                SiteWorkflowStep.step_name == WorkflowStepName.CDA_EXECUTION.value,
            )
        )

    step = step_result.scalar_one_or_none()
    if not step:
        raise HTTPException(status_code=404, detail="CDA Execution step not found")

    step_data = step.step_data or {}

    # Validate CDA is required
    if step_data.get("cda_required") is not True:
        raise HTTPException(
            status_code=400,
            detail="Cannot sign CDA: CDA is not marked as required",
        )

    # Check if already internally signed
    if step_data.get("cda_status") in ("internally_signed", "SIGNED"):
        raise HTTPException(
            status_code=400,
            detail="CDA has already been signed internally",
        )

    # Get study and site names for document generation
    site_name = site.name or "Unknown Site"
    study_name = "Unknown Study"
    if study_site:
        study_result = await db.execute(select(Study).where(Study.id == study_site.study_id))
        study_obj = study_result.scalar_one_or_none()
        if study_obj:
            study_name = study_obj.name or study_obj.study_id or "Unknown Study"
    
    # Update step_data with internal signature
    internal_signed_at = datetime.now(timezone.utc).isoformat()
    step_data.update(
        {
            "internal_signer_name": signer_name,
            "internal_signer_title": signer_title,
            "internal_signed_at": internal_signed_at,
            "cda_status": "internally_signed",
        }
    )
    
    # Generate and store CDA document HTML
    cda_template = step_data.get('cda_template', 'standard')
    html_content = generate_cda_document_html(
        study_name=study_name,
        site_name=site_name,
        cda_template=cda_template,
        step_data=step_data,
    )
    
    # Get document path (same file for this study-site combination)
    html_file_path, file_id_for_url = get_cda_document_path(
        study_site_id=study_site.id if study_site else None,
        site_id=site.id,
        study_name=study_name,
        site_name=site_name,
    )
    
    # Save HTML document
    try:
        with open(html_file_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        logger.info(f"CDA document saved successfully: {html_file_path}")
        # Verify file was created
        if not html_file_path.exists():
            logger.error(f"CDA document file was not created: {html_file_path}")
            raise Exception(f"Failed to create CDA document file: {html_file_path}")
    except Exception as e:
        logger.error(f"Error saving CDA document: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to save CDA document: {str(e)}")
    
    # Store document URL in step_data (without /api prefix since apiBase will add it)
    step_data["cda_document_url"] = f"/api/cda/document/{file_id_for_url}"
    step_data["cda_document_path"] = str(html_file_path)

    # Update step_data - ensure SQLAlchemy tracks the JSONB change
    step.step_data = step_data.copy()  # Use copy to ensure it's a new dict reference
    step.updated_at = datetime.now(timezone.utc)
    
    # Mark the JSONB column as modified so SQLAlchemy knows to update it
    from sqlalchemy.orm.attributes import flag_modified
    flag_modified(step, "step_data")
    
    await db.commit()
    await db.refresh(step)
    
    # Verify what was actually saved by re-reading from DB
    await db.refresh(step)
    saved_step_data = step.step_data or {}
    
    # Log what was saved for debugging
    logger.info(
        f"CDA internal signing saved - step_id={step.id}, study_site_id={step.study_site_id}, "
        f"site_id={step.site_id}, cda_status={step_data.get('cda_status')}, "
        f"step_data_keys={list(step_data.keys())}, "
        f"saved_cda_status={saved_step_data.get('cda_status')}, saved_keys={list(saved_step_data.keys())}"
    )
    
    # Verify the save was successful
    if saved_step_data.get('cda_status') != 'internally_signed':
        logger.error(
            f"WARNING: CDA status was not saved correctly! "
            f"Expected 'internally_signed', got {saved_step_data.get('cda_status')}. "
            f"Step data keys: {list(saved_step_data.keys())}"
        )

    # System event -> append to pinned Public Notice Board
    try:
        # Use StudySite-scoped agreement notice helper – message and event_type unchanged
        await create_agreement_notice(
            db=db,
            agreement=agreement,
            event_type="cda_signed",
            message="CDA signed internally.",
            metadata={},
        )
    except Exception as e:
        logger.error(f"Failed to append internal-sign notice to Public Notice Board: {e}", exc_info=True)

    return {
        "message": "CDA signed internally successfully",
        "step_data": step_data,
        "document_url": step_data["cda_document_url"],
    }


@router.post("/sites/{site_id}/cda/sync-zoho-status")
async def sync_zoho_sign_status(
    site_id: str,
    study_id: Optional[str] = Query(
        None, description="Study identifier for study-specific CDA (UUID or external id)"
    ),
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    [DEPRECATED] This endpoint is deprecated. CDA execution is now handled via the Agreement module.
    
    Please use the Agreement module to sync Zoho Sign status:
    - Go to the Agreement tab
    - Select the CDA Agreement
    - Use the "Sync Status from Zoho" button in the Agreement detail page
    
    This endpoint will be removed in a future version.
    """
    raise HTTPException(
        status_code=410,  # 410 Gone - indicates resource is no longer available
        detail="This endpoint is deprecated. CDA execution is now handled via the Agreement module. Please sync Zoho Sign status from the Agreement tab instead."
    )
    from app.services.zoho_sign_service import zoho_sign_service
    from sqlalchemy import text
    
    # --------------------------------------------
    # Resolve site (same logic as send endpoint)
    # --------------------------------------------
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()
    
    if not site:
        raise HTTPException(status_code=404, detail=f"Site not found: {site_id}")
    
    # --------------------------------------------
    # Resolve study & study_site (mirror send logic)
    # --------------------------------------------
    study_site: Optional[StudySite] = None
    study_obj: Optional[Study] = None
    if study_id:
        try:
            study_uuid = UUID(str(study_id))
            study_result = await db.execute(select(Study).where(Study.id == study_uuid))
            study_obj = study_result.scalar_one_or_none()
        except (ValueError, TypeError):
            study_result = await db.execute(
                select(Study).where(
                    or_(
                        Study.study_id == str(study_id),
                        Study.name == str(study_id),
                    )
                )
            )
            study_obj = study_result.scalar_one_or_none()
        
        if study_obj:
            # Use the same helper used by send_cda_for_signature so that
            # we target the exact same study_site + CDA step.
            study_site = await get_or_create_study_site(db, study_obj.id, site.id)
    
    # --------------------------------------------
    # Locate the CDA execution step
    # Use the same selection strategy as send_cda_for_signature:
    # - Prefer study-scoped when available
    # - Fall back to site-scoped when needed
    # --------------------------------------------
    study_scoped_step: Optional[SiteWorkflowStep] = None
    site_scoped_step: Optional[SiteWorkflowStep] = None
    
    if study_site:
        study_step_result = await db.execute(
            select(SiteWorkflowStep).where(
                SiteWorkflowStep.study_site_id == study_site.id,
                SiteWorkflowStep.step_name == WorkflowStepName.CDA_EXECUTION.value,
            )
        )
        study_scoped_step = study_step_result.scalar_one_or_none()
        if study_scoped_step:
            await db.refresh(study_scoped_step)
    
    site_step_result = await db.execute(
        select(SiteWorkflowStep).where(
            SiteWorkflowStep.site_id == site.id,
            SiteWorkflowStep.step_name == WorkflowStepName.CDA_EXECUTION.value,
        )
    )
    site_scoped_step = site_step_result.scalar_one_or_none()
    if site_scoped_step:
        await db.refresh(site_scoped_step)
    
    if not study_scoped_step and not site_scoped_step:
        raise HTTPException(status_code=404, detail="CDA execution step not found")
    
    # Choose step using the same preference as send (but we don't care about
    # internal signing here – we just want the step that holds zoho_sign_request_id).
    def _status_lower(step_obj: Optional[SiteWorkflowStep]) -> str:
        if not step_obj or not step_obj.step_data:
            return ""
        return str(step_obj.step_data.get("cda_status") or "").lower()
    
    study_status = _status_lower(study_scoped_step)
    site_status = _status_lower(site_scoped_step)
    
    if study_status == "internally_signed":
        step = study_scoped_step
    elif site_status == "internally_signed":
        step = site_scoped_step
    else:
        # Neither is internally signed – fall back to whichever exists,
        # preferring study-scoped for consistency with send.
        step = study_scoped_step or site_scoped_step
    
    if not step:
        # Extremely defensive; should not happen given earlier check
        raise HTTPException(status_code=404, detail="CDA execution step not found")
    
    step_data = step.step_data or {}
    zoho_request_id = step_data.get("zoho_sign_request_id")
    
    if not zoho_request_id:
        raise HTTPException(status_code=400, detail="No Zoho Sign request ID found for this CDA")
    
    logger.info(f"Syncing Zoho Sign status for request_id: {zoho_request_id}")
    
    try:
        # Fetch current status from Zoho Sign
        request_details = zoho_sign_service.get_request_details(zoho_request_id)
        requests_data = request_details.get("requests", {})

        # Zoho's response structure can vary slightly; normalise it:
        # - Sometimes "requests" is an object, sometimes a list
        # - Status fields may be named "status", "request_status", or "request_status_name"
        if isinstance(requests_data, list) and requests_data:
            requests_data = requests_data[0]

        raw_status = (
            requests_data.get("status")
            or requests_data.get("request_status")
            or requests_data.get("request_status_name")
            or requests_data.get("request_status_text")
        )
        zoho_status = (str(raw_status) if raw_status is not None else "").upper()
        
        logger.info(
            f"Zoho Sign status for {zoho_request_id}: raw={raw_status!r}, normalized={zoho_status}"
        )
        
        # Update based on Zoho status
        if zoho_status == "COMPLETED":
            # Same logic as webhook handler / send: compute names and file path
            site_name = site.name or "Unknown Site"
            study_name = "Unknown Study"
            if study_obj:
                study_name = study_obj.name or study_obj.study_id or "Unknown Study"
            
            html_file_path, _ = get_cda_document_path(
                study_site_id=step.study_site_id,
                site_id=site.id,
                study_name=study_name,
                site_name=site_name,
            )
            signed_pdf_path = html_file_path.with_suffix('.signed.pdf')
            
            try:
                zoho_sign_service.download_signed_document(zoho_request_id, signed_pdf_path)
                logger.info(f"Downloaded signed CDA PDF: {signed_pdf_path}")
                
                # Create SiteDocument
                signed_document_id = None
                if signed_pdf_path.exists():
                    signed_pdf_size = signed_pdf_path.stat().st_size
                    signed_cda_document = SiteDocument(
                        site_id=site.id,
                        category=DocumentCategory.SIGNED_CDA,
                        file_path=str(signed_pdf_path),
                        file_name=f"CDA_Signed_{study_name.replace(' ', '_')}_{site_name.replace(' ', '_')}_{datetime.now(timezone.utc).strftime('%Y%m%d')}.pdf",
                        content_type="application/pdf",
                        size=signed_pdf_size,
                        uploaded_by=None,
                        description=f"Signed CDA via Zoho Sign - {study_name} - {site_name}",
                    )
                    db.add(signed_cda_document)
                    await db.flush()
                    await db.refresh(signed_cda_document)
                    signed_document_id = signed_cda_document.id
                
                pdf_document_id = signed_pdf_path.stem
                step_data.update({
                    "zoho_sign_status": "completed",
                    "cda_status": "CDA_COMPLETED",
                    "cda_signed_at": datetime.now(timezone.utc).isoformat(),
                    "cda_signed_document": str(signed_pdf_path),
                    "cda_document_url": f"/api/cda/document/{pdf_document_id}",
                    "cda_signed_pdf_url": f"/api/cda/signed-pdf/{pdf_document_id}",
                    "signed_cda_document_id": str(signed_document_id) if signed_document_id else None,
                })
            except Exception as e:
                logger.error(f"Failed to download signed document: {str(e)}")
                step_data.update({
                    "zoho_sign_status": "completed",
                    "cda_status": "CDA_COMPLETED",
                    "cda_signed_at": datetime.now(timezone.utc).isoformat(),
                    "zoho_sign_error": str(e),
                })
        
        elif zoho_status == "DECLINED":
            step_data.update({
                "zoho_sign_status": "DECLINED",
                "cda_status": "CDA_REJECTED",
                "cda_declined_at": datetime.now(timezone.utc).isoformat(),
            })
        
        elif zoho_status == "EXPIRED":
            step_data.update({
                "zoho_sign_status": "EXPIRED",
                "cda_status": "CDA_EXPIRED",
                "cda_expired_at": datetime.now(timezone.utc).isoformat(),
            })
        
        # Update step
        step.step_data = step_data.copy()
        step.updated_at = datetime.now(timezone.utc)
        from sqlalchemy.orm.attributes import flag_modified
        flag_modified(step, "step_data")
        
        await db.commit()
        
        # System event -> append only to pinned Public Notice Board
        try:
            notice_msg = f"CDA status synced from Zoho: {zoho_status}."
            await create_agreement_notice(
                db=db,
                agreement=agreement,
                event_type="cda_signed" if zoho_status == "COMPLETED" else "status_update",
                message=notice_msg,
                metadata={
                    "attachment_url": step_data.get("cda_signed_pdf_url") if zoho_status == "COMPLETED" else None,
                    "attachment_name": "Signed CDA" if zoho_status == "COMPLETED" else None,
                    "attachment_type": "document" if zoho_status == "COMPLETED" else None,
                },
            )
        except Exception as e:
            logger.error(f"Failed to append sync notice to Public Notice Board: {e}", exc_info=True)
        
        return {
            "status": "success",
            "message": f"Synced Zoho Sign status: {zoho_status}",
            "zoho_status": zoho_status,
            "cda_status": step_data.get("cda_status"),
        }
    
    except Exception as e:
        logger.error(f"Failed to sync Zoho Sign status: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to sync status: {str(e)}")


@router.post("/sites/{site_id}/cda/send")
async def send_cda_for_signature(
    site_id: str,
    study_id: Optional[str] = Query(
        None, description="Study identifier for study-specific CDA (UUID or external id)"
    ),
    site_signer_email: str = Form(
        default="labeshg@dizzaroo.com",
        description="Site signatory email (signs first)",
    ),
    sponsor_signer_email: str = Form(
        default="labeshg@dizzaroo.com",
        description="Sponsor/Internal signatory email (signs second)",
    ),
    cc_emails: Optional[str] = Form(
        default="labeshg@dizzaroo.com",
        description="CC recipients (comma-separated)",
    ),
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    [DEPRECATED] This endpoint is deprecated. CDA execution is now handled via the Agreement module.
    
    Please use the Agreement module to create and execute CDA Agreements:
    - Create a CDA Agreement in the Agreement tab
    - Use the "Send for Signature" button in the Agreement detail page
    
    This endpoint will be removed in a future version.
    """
    raise HTTPException(
        status_code=410,  # 410 Gone - indicates resource is no longer available
        detail="This endpoint is deprecated. CDA execution is now handled via the Agreement module. Please create and execute a CDA Agreement in the Agreement tab instead."
    )
    from app.services.zoho_sign_service import zoho_sign_service
    from app.utils.pdf_generator import html_to_pdf

    # Resolve site
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()

    if not site:
        raise HTTPException(status_code=404, detail="Site not found")

    # Resolve study & study_site for study-specific CDA (if study_id provided)
    study_site: Optional[StudySite] = None
    study_obj: Optional[Study] = None
    if study_id:
        try:
            study_uuid = UUID(str(study_id))
            study_result = await db.execute(select(Study).where(Study.id == study_uuid))
            study_obj = study_result.scalar_one_or_none()
        except (ValueError, TypeError):
            study_result = await db.execute(
                select(Study).where(
                    or_(
                        Study.study_id == str(study_id),
                        Study.name == str(study_id),
                    )
                )
            )
            study_obj = study_result.scalar_one_or_none()

        if study_obj:
            study_site = await get_or_create_study_site(db, study_obj.id, site.id)
        # If study cannot be resolved, fall back to site-level CDA step

    # Get CDA Execution step
    #
    # IMPORTANT: We may have both a site-scoped CDA step and a study_site-scoped CDA step.
    # Internal signing is study-scoped when study_id is provided; however, some environments
    # may still have site-level CDA steps created earlier. To avoid mismatches (signed one
    # step but trying to send from the other), we look up BOTH and pick the one that is
    # actually internally signed.
    study_scoped_step: Optional[SiteWorkflowStep] = None
    site_scoped_step: Optional[SiteWorkflowStep] = None

    if study_site:
        study_step_result = await db.execute(
            select(SiteWorkflowStep).where(
                SiteWorkflowStep.study_site_id == study_site.id,
                SiteWorkflowStep.step_name == WorkflowStepName.CDA_EXECUTION.value,
            )
        )
        study_scoped_step = study_step_result.scalar_one_or_none()
        # Refresh to ensure we have latest data from DB
        if study_scoped_step:
            await db.refresh(study_scoped_step)

    site_step_result = await db.execute(
        select(SiteWorkflowStep).where(
            SiteWorkflowStep.site_id == site.id,
            SiteWorkflowStep.step_name == WorkflowStepName.CDA_EXECUTION.value,
        )
    )
    site_scoped_step = site_step_result.scalar_one_or_none()
    # Refresh to ensure we have latest data from DB
    if site_scoped_step:
        await db.refresh(site_scoped_step)

    if not study_scoped_step and not site_scoped_step:
        raise HTTPException(status_code=404, detail="CDA Execution step not found")

    def _status_lower(step_obj: Optional[SiteWorkflowStep]) -> str:
        if not step_obj or not step_obj.step_data:
            return ""
        return str(step_obj.step_data.get("cda_status") or "").lower()
    
    def _log_step_details(step_obj: Optional[SiteWorkflowStep], label: str):
        """Helper to log step details for debugging"""
        if step_obj:
            step_data_keys = list(step_obj.step_data.keys()) if step_obj.step_data else []
            cda_status = step_obj.step_data.get("cda_status") if step_obj.step_data else None
            logger.info(
                f"{label} - step_id={step_obj.id}, study_site_id={step_obj.study_site_id}, "
                f"site_id={step_obj.site_id}, cda_status={cda_status}, keys={step_data_keys}"
            )
        else:
            logger.info(f"{label} - step not found")

    # Log both steps for debugging
    _log_step_details(study_scoped_step, "Study-scoped step")
    _log_step_details(site_scoped_step, "Site-scoped step")
    
    # Also log the study_site lookup for debugging
    if study_site:
        logger.info(f"Looking up CDA step for study_site_id={study_site.id}, site_id={site.id}")
    else:
        logger.info(f"Looking up CDA step for site_id={site.id} (no study_site)")

    # Prefer the study-scoped step if it's internally signed; otherwise fall back to whichever is signed
    study_status = _status_lower(study_scoped_step)
    site_status = _status_lower(site_scoped_step)
    
    if study_status == "internally_signed":
        step = study_scoped_step
        logger.info(
            f"Sending CDA using study-scoped step (study_site_id={getattr(study_site, 'id', None)}, step_id={step.id})"
        )
    elif site_status == "internally_signed":
        step = site_scoped_step
        logger.info(f"Sending CDA using site-scoped step (site_id={site.id}, step_id={step.id})")
    else:
        # Neither step is internally signed; default to study-scoped if it exists (better UX)
        step = study_scoped_step or site_scoped_step
        logger.warning(
            "Sending CDA fallback step selection (not internally signed yet): "
            f"study_scoped_status={study_status!r}, "
            f"site_scoped_status={site_status!r}, "
            f"chosen_step_id={step.id if step else None}, "
            f"study_site_id={getattr(study_site, 'id', None)}"
        )

    step_data = step.step_data or {}

    # Validate CDA is required and internally signed
    cda_required = step_data.get("cda_required")
    cda_status = step_data.get("cda_status")
    
    # Log for debugging
    logger.info(f"CDA send validation - cda_required: {cda_required} (type: {type(cda_required)}), cda_status: {cda_status}, step_data keys: {list(step_data.keys())}")
    
    # Check cda_required - accept True (bool) or truthy string values
    if cda_required is not True and str(cda_required).lower() not in ("true", "yes", "1"):
        raise HTTPException(
            status_code=400,
            detail=f"Cannot send CDA: CDA is not marked as required (current value: {cda_required}, type: {type(cda_required)})",
        )

    # No need to check for internal signing - we send directly to Zoho Sign
    # Just ensure CDA is required

    # Get study and site names for document generation
    site_name = site.name or "Unknown Site"
    study_name = "Unknown Study"
    if study_obj:
        study_name = study_obj.name or study_obj.study_id or "Unknown Study"
    
    # Generate CDA HTML document (without internal signature - will be signed via Zoho Sign)
    cda_template = step_data.get('cda_template', 'standard')
    # Create step_data without internal signature for document generation
    doc_step_data = step_data.copy()
    # Don't include internal signature since we're using Zoho Sign
    html_content = generate_cda_document_html(
        study_name=study_name,
        site_name=site_name,
        cda_template=cda_template,
        step_data=doc_step_data,
    )
    
    # Get document path for PDF
    html_file_path, file_id_for_url = get_cda_document_path(
        study_site_id=study_site.id if study_site else None,
        site_id=site.id,
        study_name=study_name,
        site_name=site_name,
    )
    
    # Generate PDF from HTML
    pdf_path = html_file_path.with_suffix('.pdf')
    try:
        html_to_pdf(html_content, pdf_path)
        logger.info(f"Generated CDA PDF: {pdf_path}")
    except Exception as e:
        logger.error(f"Failed to generate PDF: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate PDF: {str(e)}",
        )
    
    # Prepare Zoho Sign recipients with signing order
    # Site signs first (order=1), Sponsor signs second (order=2)
    recipients = [
        {
            "email": site_signer_email.strip(),
            "name": "Site Signatory",
            "action": "SIGN",
            "signing_order": 1,
        },
        {
            "email": sponsor_signer_email.strip(),
            "name": "Sponsor Signatory",
            "action": "SIGN",
            "signing_order": 2,
        },
    ]
    
    # Parse CC emails
    cc_list = []
    if cc_emails:
        cc_list = [email.strip() for email in cc_emails.split(",") if email.strip()]
    
    # Create Zoho Sign request
    request_name = f"CDA - {study_name} - {site_name}"
    message = f"Please review and sign the Confidential Disclosure Agreement for {study_name} at {site_name}."
    
    try:
        zoho_response = zoho_sign_service.create_signature_request(
            request_name=request_name,
            document_path=pdf_path,
            recipients=recipients,
            cc_recipients=cc_list if cc_list else None,
            message=message,
        )
        
        # Extract request ID from response
        requests_data = zoho_response.get("requests", {})
        zoho_request_id = requests_data.get("request_id")
        
        if not zoho_request_id:
            logger.error(f"Zoho Sign response missing request_id: {zoho_response}")
            raise HTTPException(
                status_code=500,
                detail="Zoho Sign request created but no request_id returned",
            )
        
        logger.info(f"Created Zoho Sign request: {zoho_request_id}")
        
    except Exception as e:
        logger.error(f"Failed to create Zoho Sign request: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create Zoho Sign request: {str(e)}",
        )
    
    # Update step_data with Zoho Sign metadata
    step_data.update(
        {
            "zoho_sign_request_id": zoho_request_id,
            "zoho_sign_status": "SENT",
            "cda_status": "CDA_SENT",  # New status for Zoho Sign integration
            "cda_sent_at": datetime.now(timezone.utc).isoformat(),
            "cda_sent_by": current_user.get("user_id") if current_user else None,
            "site_signer_email": site_signer_email.strip(),
            "sponsor_signer_email": sponsor_signer_email.strip(),
            "cc_emails": cc_list,
        }
    )

    # Update step_data - ensure SQLAlchemy tracks the JSONB change
    step.step_data = step_data.copy()
    step.updated_at = datetime.now(timezone.utc)
    
    # Mark the JSONB column as modified so SQLAlchemy knows to update it
    from sqlalchemy.orm.attributes import flag_modified
    flag_modified(step, "step_data")
    
    await db.commit()
    await db.refresh(step)
    
    logger.info(
        f"CDA sent via Zoho Sign - step_id={step.id}, zoho_request_id={zoho_request_id}, "
        f"site_signer={site_signer_email}, sponsor_signer={sponsor_signer_email}"
    )

    # System event -> append only to pinned Public Notice Board.
    try:
        await create_agreement_notice(
            db=db,
            agreement=agreement,
            event_type="cda_sent",
            message=f"CDA sent via Zoho Sign for site {site.name or site.site_id}.",
            metadata={},
        )
    except Exception as e:
        logger.error(f"Failed to append CDA-send notice to Public Notice Board: {e}", exc_info=True)

    return {
        "message": "CDA sent for signature via Zoho Sign successfully",
        "zoho_request_id": zoho_request_id,
        "status": "CDA_SENT",
        "site_signer_email": site_signer_email.strip(),
        "sponsor_signer_email": sponsor_signer_email.strip(),
    }


async def _handle_agreement_webhook(
    db: AsyncSession,
    agreement: Agreement,
    event: str,
    zoho_request_id: str,
    requests_data: dict,
):
    """
    Handle Zoho Sign webhook for agreements.
    
    Events:
    - request.completed: Download signed PDF, create AgreementSignedDocument, update status to EXECUTED
    - request.declined: Update signature_status to DECLINED
    - request.expired: Update signature_status to EXPIRED
    """
    from app.services.zoho_sign_service import zoho_sign_service
    from pathlib import Path
    
    # Reload agreement with documents
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement.id)
        .options(
            selectinload(Agreement.documents),
            selectinload(Agreement.signed_documents)
        )
    )
    agreement = agreement_result.scalar_one()
    
    if event == "request.completed":
        # Download signed PDF from Zoho Sign
        try:
            # Determine signed PDF path
            signed_dir = Path(settings.upload_dir) / "signed_agreements"
            signed_dir.mkdir(parents=True, exist_ok=True)
            signed_pdf_path = signed_dir / f"agreement_{agreement.id}_signed_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.pdf"
            
            # Download signed PDF
            zoho_sign_service.download_signed_document(zoho_request_id, signed_pdf_path)
            logger.info(f"Downloaded signed agreement PDF: {signed_pdf_path}")
            
            # Get signed_at from Zoho response (if available)
            signed_at = None
            if requests_data.get("completed_time"):
                try:
                    signed_at = datetime.fromisoformat(requests_data["completed_time"].replace("Z", "+00:00"))
                except:
                    pass
            
            # Create AgreementSignedDocument record
            signed_document = AgreementSignedDocument(
                agreement_id=agreement.id,
                file_path=str(signed_pdf_path),
                signed_at=signed_at,
                zoho_request_id=zoho_request_id,
            )
            db.add(signed_document)
            await db.flush()
            
            # Get latest document version number
            max_version_result = await db.execute(
                select(func.max(AgreementDocument.version_number))
                .where(AgreementDocument.agreement_id == agreement.id)
            )
            max_version_number = max_version_result.scalar_one_or_none()
            next_version_number = (max_version_number if max_version_number else 0) + 1
            
            # Get latest document content
            latest_document = max(agreement.documents, key=lambda d: d.version_number)
            
            # Create new AgreementDocument version marked as signed
            signed_document_version = AgreementDocument(
                agreement_id=agreement.id,
                version_number=next_version_number,
                document_content=latest_document.document_content,  # Same content as last editable version
                document_html='',  # Legacy field
                created_from_template_id=None,
                created_by=None,  # System-created
                is_signed_version='true',
            )
            db.add(signed_document_version)
            await db.flush()
            
            # Update agreement status
            agreement.status = AgreementStatusEnum.EXECUTED
            agreement.signature_status = "COMPLETED"
            
            # Create SYSTEM comment
            # Note: version_id is for legacy file-based versions only, not AgreementDocument
            # For document-based agreements, we pass None
            await create_agreement_system_comment(
                db,
                agreement.id,
                None,  # version_id is only for legacy file-based versions, not AgreementDocument
                "Agreement executed. Signed document received from Zoho Sign."
            )
            
            # Create Notice Board entry for EXECUTED status
            try:
                await create_agreement_notice_board_entry(
                    db,
                    agreement,
                    AgreementStatusEnum.EXECUTED.value,
                    created_by=None  # System event
                )
            except Exception as e:
                logger.error(f"Failed to create Notice Board entry for agreement {agreement.id}: {str(e)}", exc_info=True)
                # Don't fail the transaction if notice creation fails
            
            # PART 2: Update Site Status Milestone - Complete CDA Execution
            # Only for CDA agreements
            # Ensure documents are loaded for type checking
            if not hasattr(agreement, 'documents') or agreement.documents is None:
                await db.refresh(agreement, ['documents'])
            
            try:
                await complete_cda_execution_milestone(
                    db,
                    agreement,
                    created_by=None  # System event
                )
            except Exception as e:
                logger.error(f"Failed to complete CDA Execution milestone for agreement {agreement.id}: {str(e)}", exc_info=True)
                # Don't fail the transaction if milestone update fails
            
            await db.commit()
            logger.info(f"Agreement {agreement.id} marked as EXECUTED after signature completion")
            
        except Exception as e:
            logger.error(f"Failed to process completed signature for agreement {agreement.id}: {str(e)}")
            # Still update status even if download fails
            agreement.signature_status = "COMPLETED"
            agreement.status = AgreementStatusEnum.EXECUTED
            await create_agreement_system_comment(
                db,
                agreement.id,
                None,
                f"Agreement signature completed, but PDF download failed: {str(e)}"
            )
            await db.commit()
    
    elif event == "request.declined":
        agreement.signature_status = "DECLINED"
        await create_agreement_system_comment(
            db,
            agreement.id,
            None,
            "Agreement signature request was declined."
        )
        await db.commit()
        logger.info(f"Agreement {agreement.id} signature declined")
    
    elif event == "request.expired":
        agreement.signature_status = "EXPIRED"
        await create_agreement_system_comment(
            db,
            agreement.id,
            None,
            "Agreement signature request expired."
        )
        await db.commit()
        logger.info(f"Agreement {agreement.id} signature expired")
    
    else:
        logger.warning(f"Unhandled agreement webhook event: {event}")


@router.post("/webhooks/zoho-sign")
async def zoho_sign_webhook(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """
    Webhook endpoint to handle Zoho Sign events.
    Handles: request.completed, request.declined, request.expired
    
    Expected payload structure:
    {
        "event": "request.completed" | "request.declined" | "request.expired",
        "requests": {
            "request_id": "...",
            "request_name": "...",
            "status": "...",
            ...
        }
    }
    """
    from app.services.zoho_sign_service import zoho_sign_service
    from app.utils.pdf_generator import html_to_pdf
    from sqlalchemy import text
    
    request_data = await request.json()
    event = request_data.get("event")
    requests_data = request_data.get("requests", {})
    zoho_request_id = requests_data.get("request_id")
    
    if not zoho_request_id:
        logger.warning(f"Zoho Sign webhook missing request_id: {request_data}")
        raise HTTPException(status_code=400, detail="Missing request_id in webhook payload")
    
    logger.info(f"Zoho Sign webhook received - event: {event}, request_id: {zoho_request_id}")
    
    # Try to find Agreement first (newer workflow)
    agreement_result = await db.execute(
        select(Agreement).where(Agreement.zoho_request_id == zoho_request_id)
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if agreement:
        # Handle agreement signature webhook
        await _handle_agreement_webhook(db, agreement, event, zoho_request_id, requests_data)
        return {"status": "acknowledged", "message": "Agreement webhook processed"}
    
    # Fall back to CDA step lookup (existing workflow)
    result = await db.execute(
        text("""
            SELECT sws.id, sws.step_data, sws.study_site_id, sws.site_id
            FROM site_workflow_steps sws
            WHERE sws.step_name = 'cda_execution'
              AND sws.step_data->>'zoho_sign_request_id' = :request_id
        """),
        {"request_id": zoho_request_id},
    )
    
    row = result.fetchone()
    if not row:
        logger.warning(f"Neither agreement nor CDA step found for Zoho request_id: {zoho_request_id}")
        # Return 200 to acknowledge webhook even if we can't find the step
        return {"status": "acknowledged", "message": "Request ID not found in system"}
    
    step_id = row[0]
    step_data = row[1] or {}
    study_site_id = row[2]
    site_id = row[3]
    
    # Get step object for update
    step_result = await db.execute(
        select(SiteWorkflowStep).where(SiteWorkflowStep.id == step_id)
    )
    step = step_result.scalar_one_or_none()
    
    if not step:
        logger.error(f"Step not found for step_id: {step_id}")
        return {"status": "error", "message": "Step not found"}
    
    # Handle different events
    if event == "request.completed":
        # Download signed PDF from Zoho Sign
        try:
            # Get site and study info for file path
            site_result = await db.execute(select(Site).where(Site.id == site_id))
            site = site_result.scalar_one_or_none()
            site_name = site.name if site else "Unknown Site"
            
            study_name = "Unknown Study"
            if study_site_id:
                study_site_result = await db.execute(
                    select(StudySite).where(StudySite.id == study_site_id)
                )
                study_site_obj = study_site_result.scalar_one_or_none()
                if study_site_obj:
                    study_result = await db.execute(
                        select(Study).where(Study.id == study_site_obj.study_id)
                    )
                    study_obj = study_result.scalar_one_or_none()
                    if study_obj:
                        study_name = study_obj.name or study_obj.study_id or "Unknown Study"
            
            # Determine signed PDF path
            html_file_path, _ = get_cda_document_path(
                study_site_id=study_site_id,
                site_id=site_id,
                study_name=study_name,
                site_name=site_name,
            )
            signed_pdf_path = html_file_path.with_suffix('.signed.pdf')
            
            # Download signed PDF
            zoho_sign_service.download_signed_document(zoho_request_id, signed_pdf_path)
            logger.info(f"Downloaded signed CDA PDF: {signed_pdf_path}")
            
            # Create SiteDocument record for the signed PDF
            signed_document_id = None
            if site_id and signed_pdf_path.exists():
                signed_pdf_size = signed_pdf_path.stat().st_size
                signed_cda_document = SiteDocument(
                    site_id=site_id,
                    category=DocumentCategory.SIGNED_CDA,
                    file_path=str(signed_pdf_path),
                    file_name=f"CDA_Signed_{study_name.replace(' ', '_')}_{site_name.replace(' ', '_')}_{datetime.now(timezone.utc).strftime('%Y%m%d')}.pdf",
                    content_type="application/pdf",
                    size=signed_pdf_size,
                    uploaded_by=None,  # External signer via Zoho Sign
                    description=f"Signed CDA via Zoho Sign - {study_name} - {site_name}",
                )
                db.add(signed_cda_document)
                await db.flush()
                await db.refresh(signed_cda_document)
                signed_document_id = signed_cda_document.id
                logger.info(f"Created SiteDocument for signed CDA PDF: {signed_document_id}")
            
            # Update step_data with signed PDF URL
            pdf_document_id = signed_pdf_path.stem
            step_data.update({
                "zoho_sign_status": "completed",  # Use lowercase to match user's requirement
                "cda_status": "CDA_COMPLETED",
                "cda_signed_at": datetime.now(timezone.utc).isoformat(),
                "cda_signed_document": str(signed_pdf_path),
                "cda_document_url": f"/api/cda/document/{pdf_document_id}",
                "cda_signed_pdf_url": f"/api/cda/signed-pdf/{pdf_document_id}",  # Direct PDF endpoint
                "signed_cda_document_id": str(signed_document_id) if signed_document_id else None,
            })
            
        except Exception as e:
            logger.error(f"Failed to download signed document: {str(e)}")
            # Still update status even if download fails
            step_data.update({
                "zoho_sign_status": "COMPLETED",
                "cda_status": "CDA_COMPLETED",
                "cda_signed_at": datetime.now(timezone.utc).isoformat(),
                "zoho_sign_error": str(e),
            })
    
    elif event == "request.declined":
        step_data.update({
            "zoho_sign_status": "DECLINED",
            "cda_status": "CDA_REJECTED",
            "cda_declined_at": datetime.now(timezone.utc).isoformat(),
        })
        logger.info(f"CDA request declined: {zoho_request_id}")
    
    elif event == "request.expired":
        step_data.update({
            "zoho_sign_status": "EXPIRED",
            "cda_status": "CDA_EXPIRED",
            "cda_expired_at": datetime.now(timezone.utc).isoformat(),
        })
        logger.info(f"CDA request expired: {zoho_request_id}")
    
    else:
        logger.warning(f"Unhandled Zoho Sign webhook event: {event}")
        return {"status": "acknowledged", "message": f"Event {event} not handled"}
    
    # Update step
    step.step_data = step_data.copy()
    step.updated_at = datetime.now(timezone.utc)
    from sqlalchemy.orm.attributes import flag_modified
    flag_modified(step, "step_data")
    
    await db.commit()
    
    logger.info(f"Updated CDA step {step_id} - event: {event}, status: {step_data.get('cda_status')}")
    
    # System event -> append only to pinned Public Notice Board.
    try:
        from app.utils.system_notices import create_system_notice_message
        site_obj = None
        study_obj = None
        if site_id:
            site_result = await db.execute(select(Site).where(Site.id == site_id))
            site_obj = site_result.scalar_one_or_none()
        if study_site_id:
            study_site_result = await db.execute(select(StudySite).where(StudySite.id == study_site_id))
            study_site = study_site_result.scalar_one_or_none()
            if study_site and study_site.study_id:
                study_result = await db.execute(select(Study).where(Study.id == study_site.study_id))
                study_obj = study_result.scalar_one_or_none()

        site_id_str = site_obj.site_id if site_obj and hasattr(site_obj, "site_id") else (str(site_id) if site_id else None)
        study_id_str = study_obj.study_id if study_obj and hasattr(study_obj, "study_id") else None
        notice_message = {
            "request.completed": "CDA signed by all parties.",
            "request.declined": "CDA was declined.",
            "request.expired": "CDA request expired.",
        }.get(event, f"CDA status updated: {event}")
        
        # Map event to event_type
        event_type_map = {
            "request.completed": "cda_signed",
            "request.declined": "status_update",
            "request.expired": "status_update",
        }
        event_type = event_type_map.get(event, "status_update")

        if site_id_str:
            await create_system_notice_message(
                db=db,
                site_id=site_id_str,
                study_id=study_id_str,
                message=notice_message,
                created_by=None,
                attachment_url=step_data.get("cda_signed_pdf_url") if event == "request.completed" else None,
                attachment_name="Signed CDA" if event == "request.completed" else None,
                attachment_type="document" if event == "request.completed" else None,
                event_type=event_type,
            )
    except Exception as e:
        logger.error(f"Failed to append webhook notice to Public Notice Board: {e}", exc_info=True)
    
    return {
        "status": "success",
        "message": f"Webhook processed for event: {event}",
        "step_id": str(step_id),
        "cda_status": step_data.get("cda_status"),
    }


@router.get("/cda/sign")
async def get_cda_sign_page(
    token: str = Query(...),
    db: AsyncSession = Depends(get_db),
):
    """
    Public endpoint to get CDA signing page data.
    No authentication required - uses token for access.
    """
    from sqlalchemy import text

    # Find CDA step by token using JSONB query
    # Log the token for debugging
    logger.info(f"Looking up CDA step with token: {token[:10]}...")
    
    result = await db.execute(
        text(
            """
            SELECT sws.id, sws.step_data, ss.id AS study_site_id, s.id AS site_id, s.name AS site_name,
                   st.id AS study_id, st.name AS study_name
            FROM site_workflow_steps sws
            LEFT JOIN study_sites ss ON sws.study_site_id = ss.id
            LEFT JOIN sites s ON sws.site_id = s.id OR ss.site_id = s.id
            LEFT JOIN studies st ON ss.study_id = st.id
            WHERE sws.step_name = 'cda_execution'
              AND sws.step_data->>'cda_sign_token' = :token
            """
        ),
        {"token": token},
    )

    row = result.fetchone()
    if not row:
        # Log all CDA steps with tokens for debugging
        debug_result = await db.execute(
            text(
                """
                SELECT sws.id, sws.step_data->>'cda_sign_token' AS token, sws.study_site_id, sws.site_id
                FROM site_workflow_steps sws
                WHERE sws.step_name = 'cda_execution'
                  AND sws.step_data->>'cda_sign_token' IS NOT NULL
                """
            )
        )
        debug_rows = debug_result.fetchall()
        logger.warning(
            f"Token not found. Looking for: {token[:10]}... "
            f"Found {len(debug_rows)} CDA steps with tokens: {[(str(r[0]), r[1][:10] + '...' if r[1] else None) for r in debug_rows]}"
        )
        raise HTTPException(status_code=404, detail="Invalid or expired token")

    # Get step_id and step_data
    step_id = row[0]
    step_data = row[1] or {}

    # Check token expiration
    expires_at_str = step_data.get("cda_sign_token_expires_at")
    if expires_at_str:
        expires_at = datetime.fromisoformat(expires_at_str.replace("Z", "+00:00"))
        if expires_at < datetime.now(timezone.utc):
            raise HTTPException(status_code=400, detail="Token has expired")

    # Check if already signed
    if step_data.get("cda_status") == "SIGNED":
        raise HTTPException(status_code=400, detail="CDA has already been signed")

    # Internal signer info
    internal_signer_name = step_data.get("internal_signer_name", "")
    internal_signer_title = step_data.get("internal_signer_title", "")
    internal_signed_at = step_data.get("internal_signed_at", "")

    site_name = row[4] if len(row) > 4 else "Unknown Site"
    study_name = row[6] if len(row) > 6 else "Unknown Study"

    # Get document URL if available
    document_url = step_data.get("cda_document_url") or step_data.get("final_document_url")
    
    return {
        "token": token,
        "study_name": study_name,
        "site_name": site_name,
        "internal_signer_name": internal_signer_name,
        "internal_signer_title": internal_signer_title,
        "internal_signed_at": internal_signed_at,
        "cda_template": step_data.get("cda_template", "standard"),
        "cda_document_url": document_url,
    }


@router.post("/cda/sign")
async def sign_cda_external(
    token: str = Form(...),
    site_signer_name: str = Form(...),
    site_signer_title: str = Form(...),
    db: AsyncSession = Depends(get_db),
):
    """
    Public endpoint for external (hospital) signing of CDA.
    No authentication required - uses token for access.
    Creates HTML snapshot of signed CDA.
    """
    from sqlalchemy import text

    # Find CDA step by token
    result = await db.execute(
        text(
            """
            SELECT sws.id, sws.step_data, ss.id AS study_site_id, s.id AS site_id, s.name AS site_name,
                   st.id AS study_id, st.name AS study_name
            FROM site_workflow_steps sws
            LEFT JOIN study_sites ss ON sws.study_site_id = ss.id
            LEFT JOIN sites s ON sws.site_id = s.id OR ss.site_id = s.id
            LEFT JOIN studies st ON ss.study_id = st.id
            WHERE sws.step_name = 'cda_execution'
              AND sws.step_data->>'cda_sign_token' = :token
            """
        ),
        {"token": token},
    )

    row = result.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Invalid or expired token")

    step_id = row[0]
    step_result = await db.execute(
        select(SiteWorkflowStep).where(SiteWorkflowStep.id == step_id)
    )
    step = step_result.scalar_one_or_none()
    if not step:
        raise HTTPException(status_code=404, detail="CDA step not found")

    step_data = step.step_data or {}

    # Check token expiration
    expires_at_str = step_data.get("cda_sign_token_expires_at")
    if expires_at_str:
        expires_at = datetime.fromisoformat(expires_at_str.replace("Z", "+00:00"))
        if expires_at < datetime.now(timezone.utc):
            raise HTTPException(status_code=400, detail="Token has expired")

    # Check if already signed
    if step_data.get("cda_status") == "SIGNED":
        raise HTTPException(status_code=400, detail="CDA has already been signed")

    # Validate internal signature exists
    if step_data.get("cda_status") != "internally_signed":
        raise HTTPException(
            status_code=400,
            detail="CDA must be signed internally before external signing",
        )

    site_name = row[4] if len(row) > 4 else "Unknown Site"
    study_name = row[6] if len(row) > 6 else "Unknown Study"
    site_id = row[3] if len(row) > 3 else None
    study_site_id = row[2] if len(row) > 2 else None

    # Get site UUID if we have site_id
    site_uuid = None
    if site_id:
        try:
            site_uuid = UUID(str(site_id))
        except (ValueError, TypeError):
            site_result = await db.execute(select(Site).where(Site.site_id == str(site_id)))
            site_obj = site_result.scalar_one_or_none()
            if site_obj:
                site_uuid = site_obj.id

    # Update the existing CDA document (same file, not create new)
    now = datetime.now(timezone.utc)
    site_signed_at_str = now.isoformat()
    
    # Get existing document path or create new one
    existing_doc_path = step_data.get('cda_document_path')
    if existing_doc_path and Path(existing_doc_path).exists():
        # Update existing document
        html_file_path = Path(existing_doc_path)
        file_id_for_url = html_file_path.stem
    else:
        # Create new document path (shouldn't happen if internal signing worked, but handle gracefully)
        if site_uuid:
            html_file_path, file_id_for_url = get_cda_document_path(
                study_site_id=study_site_id,
                site_id=site_uuid,
                study_name=study_name,
                site_name=site_name,
            )
        else:
            # Fallback: use random UUID if site not found
            upload_dir = Path(settings.upload_dir)
            if not upload_dir.is_absolute():
                import pathlib
                app_dir = pathlib.Path(__file__).parent.parent
                upload_dir = app_dir / upload_dir
            cda_dir = upload_dir / "cda_snapshots"
            cda_dir.mkdir(parents=True, exist_ok=True)
            file_id = uuid.uuid4()
            html_file_path = cda_dir / f"{file_id}.html"
            file_id_for_url = str(file_id)
    
    # Generate updated HTML with both signatures
    cda_template = step_data.get('cda_template', 'standard')
    html_content = generate_cda_document_html(
        study_name=study_name,
        site_name=site_name,
        cda_template=cda_template,
        step_data=step_data,
        site_signer_name=site_signer_name,
        site_signer_title=site_signer_title,
        site_signed_at=site_signed_at_str,
    )

    # Save/update HTML document
    with open(html_file_path, "w", encoding="utf-8") as f:
        f.write(html_content)

    # Create SiteDocument record for the signed CDA
    if not site_id and study_site_id:
        study_site_result = await db.execute(
            select(StudySite).where(StudySite.id == study_site_id)
        )
        study_site_obj = study_site_result.scalar_one_or_none()
        if study_site_obj:
            site_id = study_site_obj.site_id

    document_id = None
    if site_id:
        signed_cda_document = SiteDocument(
            site_id=site_id,
            category=DocumentCategory.SIGNED_CDA,
            file_path=str(html_file_path),
            file_name=f"CDA_Signed_{study_name.replace(' ', '_')}_{site_name.replace(' ', '_')}_{now.strftime('%Y%m%d')}.html",
            content_type="text/html",
            size=len(html_content.encode("utf-8")),
            uploaded_by=None,  # External signer
            description=f"Signed CDA - {study_name} - {site_name}",
        )
        db.add(signed_cda_document)
        await db.flush()
        await db.refresh(signed_cda_document)
        document_id = signed_cda_document.id

    # Update step_data with external signature and finalize
    step_data.update(
        {
            "site_signer_name": site_signer_name,
            "site_signer_title": site_signer_title,
            "site_signed_at": site_signed_at_str,
            "cda_status": "SIGNED",  # Uppercase for existing validation
            "final_document_url": f"/api/cda/document/{file_id_for_url}",
            "cda_document_url": f"/api/cda/document/{file_id_for_url}",  # Keep consistent
            "cda_document_path": str(html_file_path),  # Update path
            "signed_cda_document_id": str(document_id) if document_id else None,
            "cda_sign_token": None,  # Invalidate token after signing
            "cda_sign_token_expires_at": None,  # Invalidate token expiration
        }
    )

    # Persist JSONB changes reliably (SQLAlchemy doesn't always detect nested dict updates)
    step.step_data = step_data.copy()
    step.updated_at = datetime.now(timezone.utc)
    from sqlalchemy.orm.attributes import flag_modified
    flag_modified(step, "step_data")
    # Do NOT auto-complete the step - internal user must complete it manually
    # This allows internal user to review and then complete the step
    # step.status = StepStatus.COMPLETED
    # step.completed_at = now
    # step.completed_by = f"External Signer: {site_signer_name}"

    await db.commit()
    await db.refresh(step)

    # Broadcast WebSocket event for real-time updates
    try:
        from app.websocket_manager import manager
        # Get site_id for broadcasting
        broadcast_site_id = str(site_uuid) if site_uuid else str(site_id) if site_id else None
        if broadcast_site_id:
            await manager.broadcast_json(
                {
                    "type": "cda_signed",
                    "site_id": broadcast_site_id,
                    "study_site_id": str(study_site_id) if study_site_id else None,
                    "cda_status": "SIGNED",
                    "step_id": str(step.id),
                    "message": "CDA has been signed externally",
                }
            )
    except Exception as e:
        logger.warning(f"Failed to broadcast CDA signed event: {e}")

    return {
        "message": "CDA signed successfully",
        "document_url": f"/api/cda/document/{file_id_for_url}",
        "signed_at": site_signed_at_str,
    }


@router.get("/cda/document/{document_id:path}")
async def get_cda_document(
    document_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    Get the CDA HTML document by document ID (filename or UUID).
    Supports both:
    - New format: filename-based (e.g., "{study_site_id}_{study}_{site}")
    - Legacy format: UUID-based (looks up in SiteDocument table)
    
    Note: Uses :path parameter to handle filenames with spaces and special characters.
    """
    from urllib.parse import unquote
    
    # URL decode the document_id to handle spaces and special characters
    document_id = unquote(document_id)
    
    upload_dir = Path(settings.upload_dir)
    if not upload_dir.is_absolute():
        import pathlib
        app_dir = pathlib.Path(__file__).parent.parent
        upload_dir = app_dir / upload_dir
    
    cda_dir = upload_dir / "cda_snapshots"
    cda_dir.mkdir(parents=True, exist_ok=True)
    
    # Try filename-based lookup first (new format)
    html_file_path = cda_dir / f"{document_id}.html"
    
    # Debug: Log what we're looking for
    logger.info(f"CDA document lookup: document_id='{document_id}', path='{html_file_path}', exists={html_file_path.exists()}")
    
    if html_file_path.exists():
        filename = f"CDA_{document_id.replace(' ', '_')}.html"
        return FileResponse(
            path=str(html_file_path),
            media_type="text/html",
            headers={
                # Force inline rendering so the HTML displays inside iframes instead of downloading
                "Content-Disposition": f'inline; filename="{filename}"'
            },
        )
    
    # If file doesn't exist, try to find it by looking up the step_data and regenerating
    # Extract study_site_id from document_id (format: {study_site_id}_{study}_{site})
    parts = document_id.split('_', 1)  # Split on first underscore
    if len(parts) >= 1:
        study_site_id_str = parts[0]
        try:
            study_site_uuid = UUID(study_site_id_str)
            # Look up the CDA step to get step_data and regenerate the document
            from sqlalchemy import text
            result = await db.execute(
                text("""
                    SELECT sws.step_data, s.name AS site_name, st.name AS study_name
                    FROM site_workflow_steps sws
                    LEFT JOIN study_sites ss ON sws.study_site_id = ss.id
                    LEFT JOIN sites s ON ss.site_id = s.id
                    LEFT JOIN studies st ON ss.study_id = st.id
                    WHERE sws.step_name = 'cda_execution'
                      AND sws.study_site_id = :study_site_id
                """),
                {"study_site_id": study_site_uuid}
            )
            row = result.fetchone()
            if row:
                step_data = row[0] or {}
                site_name = row[1] or "Unknown Site"
                study_name = row[2] or "Unknown Study"
                
                # Regenerate the document
                cda_template = step_data.get('cda_template', 'standard')
                html_content = generate_cda_document_html(
                    study_name=study_name,
                    site_name=site_name,
                    cda_template=cda_template,
                    step_data=step_data,
                )
                
                # Save it
                with open(html_file_path, "w", encoding="utf-8") as f:
                    f.write(html_content)
                
                logger.info(f"Regenerated CDA document: {html_file_path}")
                filename = f"CDA_{document_id.replace(' ', '_')}.html"
                return FileResponse(
                    path=str(html_file_path),
                    media_type="text/html",
                    headers={
                        "Content-Disposition": f'inline; filename="{filename}"'
                    },
                )
        except (ValueError, Exception) as e:
            logger.warning(f"Could not regenerate CDA document: {e}")
    
    # If not found, list available files for debugging
    existing_files = list(cda_dir.glob("*.html"))
    logger.warning(f"CDA document not found. Looking for: {html_file_path.name}, Available files: {[f.name for f in existing_files[:5]]}")
    
    # Fallback to UUID-based lookup (legacy format via SiteDocument)
    try:
        doc_uuid = UUID(document_id)
    except ValueError:
        # Not a UUID, and file doesn't exist - return 404 with helpful message
        raise HTTPException(
            status_code=404, 
            detail=f"Document not found: {document_id}. Available files: {[f.stem for f in existing_files[:3]]}"
        )

    result = await db.execute(select(SiteDocument).where(SiteDocument.id == doc_uuid))
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    file_path = Path(document.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=str(file_path),
        media_type="text/html",
        headers={"Content-Disposition": f'inline; filename="{document.file_name}"'},
    )


@router.get("/cda/signed-pdf/{document_id:path}")
async def get_signed_cda_pdf(
    document_id: str,
):
    """
    Serve the signed CDA PDF by document_id.

    The `document_id` here is the identifier we stored in step_data as
    `cda_signed_pdf_url`, e.g. `/api/cda/signed-pdf/{pdf_document_id}`.

    For signed PDFs we use paths like:
        {base}.signed.pdf

    where `pdf_document_id` is typically `{base}.signed` (Path.stem of the file).
    So the primary lookup is:
        {document_id}.pdf  ->  {base}.signed.pdf
    """
    from urllib.parse import unquote

    # URL decode to handle spaces and special characters
    document_id = unquote(document_id)

    upload_dir = Path(settings.upload_dir)
    if not upload_dir.is_absolute():
        import pathlib
        app_dir = pathlib.Path(__file__).parent.parent
        upload_dir = app_dir / upload_dir

    cda_dir = upload_dir / "cda_snapshots"
    cda_dir.mkdir(parents=True, exist_ok=True)

    # Primary guess: {document_id}.pdf  (for stems like "{base}.signed")
    pdf_path = cda_dir / f"{document_id}.pdf"

    # Fallback: treat document_id as the full filename without extension hint
    if not pdf_path.exists():
        alt_path = cda_dir / f"{document_id}.signed.pdf"
        if alt_path.exists():
            pdf_path = alt_path

    logger.info(
        f"Signed CDA PDF lookup: document_id='{document_id}', path='{pdf_path}', exists={pdf_path.exists()}"
    )

    if not pdf_path.exists():
        existing_pdfs = [f.name for f in cda_dir.glob("*.pdf")]
        raise HTTPException(
            status_code=404,
            detail=f"Signed CDA PDF not found for id: {document_id}. Available PDFs (first 3): {existing_pdfs[:3]}",
        )

    filename = pdf_path.name
    return FileResponse(
        path=str(pdf_path),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'inline; filename="{filename}"'
        },
    )


@router.get("/sites/{site_id}/cda/preview")
async def get_cda_preview(
    site_id: str,
    study_id: Optional[str] = Query(
        None,
        description="Study identifier for study-specific CDA (UUID or external id)",
    ),
    db: AsyncSession = Depends(get_db),
):
    """
    Get the current CDA document preview (works at all stages: draft, internally signed, fully signed).
    Returns the document URL if available, or generates a preview on-the-fly.
    """
    # Resolve site
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()

    if not site:
        raise HTTPException(status_code=404, detail="Site not found")

    # Resolve study_site
    study_site: Optional[StudySite] = None
    if study_id:
        study_obj: Optional[Study] = None
        try:
            study_uuid = UUID(str(study_id))
            study_result = await db.execute(select(Study).where(Study.id == study_uuid))
            study_obj = study_result.scalar_one_or_none()
        except (ValueError, TypeError):
            study_result = await db.execute(
                select(Study).where(
                    or_(
                        Study.study_id == str(study_id),
                        Study.name == str(study_id),
                    )
                )
            )
            study_obj = study_result.scalar_one_or_none()

        if study_obj:
            study_site = await get_or_create_study_site(db, study_obj.id, site.id)

    # Get CDA Execution step
    if study_site:
        step_result = await db.execute(
            select(SiteWorkflowStep).where(
                SiteWorkflowStep.study_site_id == study_site.id,
                SiteWorkflowStep.step_name == WorkflowStepName.CDA_EXECUTION.value,
            )
        )
    else:
        step_result = await db.execute(
            select(SiteWorkflowStep).where(
                SiteWorkflowStep.site_id == site.id,
                SiteWorkflowStep.step_name == WorkflowStepName.CDA_EXECUTION.value,
            )
        )

    step = step_result.scalar_one_or_none()
    if not step:
        raise HTTPException(status_code=404, detail="CDA Execution step not found")

    step_data = step.step_data or {}
    
    # Return document URL if available
    document_url = step_data.get("cda_document_url") or step_data.get("final_document_url")
    if document_url:
        return {
            "document_url": document_url,
            "status": step_data.get("cda_status", "draft"),
            "has_internal_signature": bool(step_data.get("internal_signer_name")),
            "has_external_signature": bool(step_data.get("site_signer_name")),
        }
    
    # If no document exists yet, return None (document will be created on first signature)
    return {
        "document_url": None,
        "status": step_data.get("cda_status", "draft"),
        "has_internal_signature": False,
        "has_external_signature": False,
        "message": "CDA document will be available after internal signing",
    }


# ---------------------------------------------------------------------------
# Agreement Workflow Endpoints
# ---------------------------------------------------------------------------

async def build_agreement_response(
    db: AsyncSession,
    agreement: Agreement,
    is_internal: bool = True,
    user_id: Optional[str] = None
) -> dict:
    """
    Helper function to build agreement response with role-based filtering and can_upload_new_version.
    
    Args:
        db: Database session
        agreement: Agreement object with loaded relationships
        is_internal: Whether user is internal (default True)
        user_id: Optional user ID for permission checks
        
    Returns:
        Dictionary ready for AgreementResponse schema
    """
    # Get documents (template-based)
    documents = list(agreement.documents) if hasattr(agreement, 'documents') and agreement.documents else []
    
    # Legacy file-based version system is no longer exposed at runtime; rely solely on documents.
    effective_is_legacy = agreement.is_legacy
    
    # No new legacy file uploads allowed
    can_upload = False
    
    # Filter comments by role and sort chronologically
    filtered_comments = filter_comments_by_role(list(agreement.comments), is_internal)
    
    # Get editor permissions (for backward compatibility)
    can_edit, can_comment = get_editor_permissions(agreement, is_internal)
    
    # Get comprehensive permission flags
    permissions = get_agreement_permissions(agreement, is_internal)
    
    # Get current document version number (latest version)
    current_document_version_number = None
    if documents:
        current_document_version_number = max(doc.version_number for doc in documents)
    
    return {
        "id": agreement.id,
        "site_id": agreement.site_id,
        "title": agreement.title,
        "status": agreement.status.value,
        "created_by": agreement.created_by,
        "created_at": agreement.created_at,
        "updated_at": agreement.updated_at,
        "is_legacy": effective_is_legacy,
        "documents": [
            {
                "id": doc.id,
                "agreement_id": doc.agreement_id,
                "version_number": doc.version_number,
                "document_content": doc.document_content if hasattr(doc, 'document_content') and doc.document_content else {"type": "doc", "content": []},
                "document_file_path": doc.document_file_path if hasattr(doc, 'document_file_path') else None,
                "created_from_template_id": doc.created_from_template_id,
                "created_by": doc.created_by,
                "created_at": doc.created_at,
                "is_signed_version": doc.is_signed_version,
            }
            for doc in documents
        ],
        "comments": [
            {
                "id": comment.id,
                "agreement_id": comment.agreement_id,
                "version_id": comment.version_id,
                "comment_type": comment.comment_type.value,
                "content": comment.content,
                "created_by": comment.created_by,
                "created_at": comment.created_at,
            }
            for comment in filtered_comments
        ],
        "can_upload_new_version": can_upload,
        "can_edit": permissions["can_edit"],  # Use comprehensive permissions
        "can_comment": can_comment,
        "can_save": permissions["can_save"],
        "can_move_status": permissions["can_move_status"],
        "is_locked": permissions["is_locked"],
        "current_document_version_number": current_document_version_number,
        "zoho_request_id": agreement.zoho_request_id,
        "signature_status": agreement.signature_status,
        "signed_documents": [
            {
                "id": doc.id,
                "agreement_id": doc.agreement_id,
                "file_path": doc.file_path,
                "signed_at": doc.signed_at,
                "downloaded_from_zoho_at": doc.downloaded_from_zoho_at,
                "zoho_request_id": doc.zoho_request_id,
            }
            for doc in (agreement.signed_documents if hasattr(agreement, 'signed_documents') and agreement.signed_documents else [])
        ],
    }


@router.get("/sites/{site_id}/agreements", response_model=List[schemas.AgreementResponse])
async def list_site_agreements(
    site_id: str,
    study_id: UUID | None = Query(None, description="Study identifier for scoping agreements"),
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    List all agreements for a site.
    """
    # Resolve site
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()
    
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")

    if study_id is None:
        raise HTTPException(status_code=400, detail="study_id query parameter is required")
    
    # Resolve StudySite mapping for this Study + Site pair (required)
    study_site_result = await db.execute(
        select(StudySite)
        .where(StudySite.site_id == site.id)
        .where(StudySite.study_id == study_id)
    )
    study_site = study_site_result.scalar_one_or_none()

    if not study_site:
        logger.info(
            "No StudySite mapping found for study_id=%s, site_id=%s (site_id_external=%s) – returning empty agreement list",
            str(study_id),
            str(site.id),
            site.site_id,
        )
        return []

    # Build base query for agreements with relationships (documents only)
    base_query = select(Agreement).options(
        selectinload(Agreement.comments),
        selectinload(Agreement.documents),
        selectinload(Agreement.signed_documents),
    ).order_by(Agreement.created_at.desc())

    logger.info(
        "Listing agreements for study_id=%s, site_id=%s (site_id_external=%s) via study_site_id=%s",
        str(study_id),
        str(site.id),
        site.site_id,
        str(study_site.id),
    )
    agreements_query = base_query.where(Agreement.study_site_id == study_site.id)

    agreements_result = await db.execute(agreements_query)
    agreements = agreements_result.scalars().all()
    
    # Determine if user is internal
    user_id = current_user.get("user_id") if current_user else None
    is_internal = is_user_internal(user_id, db) if user_id else True
    
    # Build responses for all agreements (can't use list comprehension with await)
    responses = []
    for agreement in agreements:
        response = await build_agreement_response(db, agreement, is_internal, user_id)
        responses.append(response)
    return responses


async def create_agreement_system_comment(
    db: AsyncSession,
    agreement_id: UUID,
    version_id: Optional[UUID],
    content: str
):
    """Helper function to create system comments for agreement events."""
    comment = AgreementComment(
        agreement_id=agreement_id,
        version_id=version_id,
        comment_type=CommentType.SYSTEM,
        content=content,
        created_by=None  # System comments have no creator
    )
    db.add(comment)
    await db.flush()
    return comment


async def create_agreement_notice_board_entry(
    db: AsyncSession,
    agreement: Agreement,
    status: str,
    created_by: Optional[str] = None
):
    """
    Create a Notice Board entry for agreement status changes.
    Only creates notices for SENT_FOR_SIGNATURE and EXECUTED statuses.
    Includes duplicate prevention to avoid creating multiple notices for the same status.

    Delegates to create_agreement_notice with identical message and metadata.
    """
    # Only create notices for specific statuses
    if status not in [AgreementStatusEnum.SENT_FOR_SIGNATURE.value, AgreementStatusEnum.EXECUTED.value]:
        return None

    if status == AgreementStatusEnum.SENT_FOR_SIGNATURE.value:
        event_type = "agreement_sent_for_signature"
        # site_id_str is derived inside helper using the same logic; keep message text identical
        message = f"Agreement '{agreement.title}' for Site '{agreement.site_id}' has been sent for signature."
    else:
        event_type = "agreement_executed"
        message = f"Agreement '{agreement.title}' for Site '{agreement.site_id}' has been executed successfully."

    try:
        await create_agreement_notice(
            db=db,
            agreement=agreement,
            event_type=event_type,
            message=message,
            metadata={
                "agreement_id": str(agreement.id),
                "agreement_status": status,
                "agreement_title": agreement.title,
            },
        )
        logger.info("Created Notice Board entry for agreement %s, status: %s", agreement.id, status)
    except Exception as e:
        logger.error(
            "Failed to create Notice Board entry for agreement %s: %s",
            agreement.id,
            str(e),
            exc_info=True,
        )


async def complete_cda_execution_milestone(
    db: AsyncSession,
    agreement: Agreement,
    created_by: Optional[str] = None
):
    """
    Complete the CDA Execution milestone when a CDA Agreement is executed.
    Reuses existing Site Status workflow logic.
    
    Args:
        db: Database session
        agreement: Agreement object (must be CDA type)
        created_by: User who triggered the execution (optional)
    
    Returns:
        Updated workflow step or None if not updated
    """
    from app.api.v1.endpoints.clinical_workflow import get_or_create_study_site
    
    # PART 1: Identify CDA Agreement
    # Check if this is a CDA agreement by checking the template type
    # Get the template from the agreement's first document
    # Ensure documents are loaded
    if not hasattr(agreement, 'documents') or not agreement.documents or len(agreement.documents) == 0:
        # Try to load documents if not already loaded
        agreement_result = await db.execute(
            select(Agreement)
            .where(Agreement.id == agreement.id)
            .options(selectinload(Agreement.documents))
        )
        agreement = agreement_result.scalar_one_or_none()
        if not agreement or not agreement.documents or len(agreement.documents) == 0:
            logger.warning(f"Agreement {agreement.id if agreement else 'unknown'} has no documents, cannot determine type")
            return None
    
    # Get template from the first document
    first_doc = agreement.documents[0]
    if not first_doc.created_from_template_id:
        logger.warning(f"Agreement {agreement.id} document has no template_id, cannot determine type")
        return None
    
    template_result = await db.execute(
        select(StudyTemplate).where(StudyTemplate.id == first_doc.created_from_template_id)
    )
    template = template_result.scalar_one_or_none()
    
    if not template:
        logger.warning(f"Template not found for agreement {agreement.id}")
        return None
    
    # Check if template type is CDA
    template_type = template.template_type.value if hasattr(template.template_type, 'value') else str(template.template_type)
    if template_type != 'CDA':
        logger.info(f"Agreement {agreement.id} is not a CDA agreement (type: {template_type}). Skipping milestone update.")
        return None
    
    # Verify site_id exists
    if not agreement.site_id:
        logger.warning(f"Agreement {agreement.id} has no site_id, cannot update milestone")
        return None
    
    # Get site
    site_result = await db.execute(select(Site).where(Site.id == agreement.site_id))
    site = site_result.scalar_one_or_none()
    if not site:
        logger.warning(f"Site not found for agreement {agreement.id}, cannot update milestone")
        return None
    
    # CDA Execution is a study-specific milestone
    # Resolve StudySite without using deprecated sites.study_id
    study_site = None

    # Prefer Agreement.study_site_id when available
    if getattr(agreement, "study_site_id", None):
        study_site_result = await db.execute(
            select(StudySite).where(StudySite.id == agreement.study_site_id)
        )
        study_site = study_site_result.scalar_one_or_none()

    # Fallback: use Agreement.study_id with get_or_create_study_site
    if not study_site and getattr(agreement, "study_id", None):
        try:
            study_uuid = UUID(str(agreement.study_id))
        except (ValueError, TypeError):
            study_uuid = None

        if study_uuid:
            study_site = await get_or_create_study_site(db, study_uuid, site.id)

    if not study_site:
        logger.warning(f"Study site not found for agreement {agreement.id}, cannot update milestone")
        return None
    
    # PART 2: Update Site Status Milestone - Complete CDA Execution step
    # Get or create the CDA Execution workflow step
    step_result = await db.execute(
        select(SiteWorkflowStep).where(
            SiteWorkflowStep.study_site_id == study_site.id,
            SiteWorkflowStep.step_name == WorkflowStepName.CDA_EXECUTION
        )
    )
    step = step_result.scalar_one_or_none()
    
    # If step doesn't exist, create it
    if not step:
        step = SiteWorkflowStep(
            study_site_id=study_site.id,
            step_name=WorkflowStepName.CDA_EXECUTION,
            status=StepStatus.NOT_STARTED,
            step_data={}
        )
        db.add(step)
        await db.flush()
    
    # PART 4: Avoid Duplicate Completion
    # Check if already completed (prevent duplicate updates)
    if step.status == StepStatus.COMPLETED:
        logger.info(f"CDA Execution milestone already completed for agreement {agreement.id}. Skipping update.")
        return step
    
    # Update step to COMPLETED
    step.status = StepStatus.COMPLETED
    step.completed_at = datetime.now(timezone.utc)
    step.completed_by = created_by or "system"
    step.updated_at = datetime.now(timezone.utc)
    
    # Update step_data to mark as completed via Agreement
    if step.step_data is None:
        step.step_data = {}
    step.step_data['completed_via_agreement'] = True
    step.step_data['agreement_id'] = str(agreement.id)
    step.step_data['cda_status'] = 'CDA_COMPLETED'
    
    await db.flush()
    logger.info(f"Completed CDA Execution milestone for agreement {agreement.id} (site: {site.site_id if hasattr(site, 'site_id') else site.id})")
    
    return step


@router.post("/sites/{site_id}/agreements", response_model=schemas.AgreementResponse)
async def create_agreement(
    site_id: str,
    agreement_data: schemas.AgreementCreate,
    study_id: UUID | None = Query(None, description="Study identifier for agreement creation"),
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Create a new agreement for a site.
    """
    # Resolve site
    try:
        site_uuid = UUID(str(site_id))
        site_result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = site_result.scalar_one_or_none()
    except (ValueError, TypeError):
        site_result = await db.execute(select(Site).where(Site.site_id == site_id))
        site = site_result.scalar_one_or_none()
    
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")

    if study_id is None:
        raise HTTPException(status_code=400, detail="study_id query parameter is required")
    
    # Check if SiteProfile exists - required before creating CDA Agreement
    profile_result = await db.execute(
        select(SiteProfile).where(SiteProfile.site_id == site.id)
    )
    profile = profile_result.scalar_one_or_none()
    
    if not profile:
        raise HTTPException(
            status_code=400,
            detail="Please complete Site Profile before creating CDA."
        )
    
    # Validate status
    try:
        status_enum = AgreementStatus(agreement_data.status.upper())
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid status: {agreement_data.status}")
    
    # For new agreements, template_id is required (non-legacy)
    template_id = getattr(agreement_data, 'template_id', None)
    is_legacy = 'false'
    
    if not template_id:
        raise HTTPException(
            status_code=400,
            detail="Template selection is required for new agreements. Please provide template_id."
        )
    
    # Validate template exists and is active
    template_result = await db.execute(
        select(StudyTemplate)
        .where(StudyTemplate.id == template_id)
        .where(StudyTemplate.is_active == 'true')
    )
    template = template_result.scalar_one_or_none()
    
    if not template:
        raise HTTPException(
            status_code=404,
            detail="Template not found or inactive"
        )

    # Resolve or create StudySite mapping for this Study + Site pair
    study_site_result = await db.execute(
        select(StudySite)
        .where(StudySite.site_id == site.id)
        .where(StudySite.study_id == study_id)
    )
    study_site = study_site_result.scalar_one_or_none()

    if not study_site:
        study_site = StudySite(
            study_id=study_id,
            site_id=site.id,
        )
        db.add(study_site)
        await db.flush()
        logger.info(
            "Created new StudySite mapping id=%s for study=%s, site=%s (site_id=%s)",
            str(study_site.id),
            str(study_id),
            str(site.id),
            site.site_id,
        )

    # Prevent duplicate agreements per StudySite + agreement_type
    existing_agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.study_site_id == study_site.id)
        .where(Agreement.agreement_type == template.template_type)
    )
    existing_agreement = existing_agreement_result.scalar_one_or_none()

    if existing_agreement:
        logger.info(
            "Duplicate agreement prevented for study_site_id=%s, type=%s (existing_agreement_id=%s)",
            str(study_site.id),
            template.template_type.value if hasattr(template.template_type, "value") else str(template.template_type),
            str(existing_agreement.id),
        )
        raise HTTPException(
            status_code=409,
            detail=f"An agreement of type {template.template_type.value if hasattr(template.template_type, 'value') else str(template.template_type)} already exists for this Study + Site."
        )

    # Create agreement (non-legacy), now explicitly linked to StudySite
    # Safety: ensure template belongs to the same study (if template.study_id is set)
    if hasattr(template, "study_id") and template.study_id and template.study_id != study_id:
        raise HTTPException(
            status_code=400,
            detail="Template does not belong to the requested study_id"
        )

    agreement = Agreement(
        site_id=site.id,
        study_id=study_id,                       # Explicit Study + Site scoping from query
        study_site_id=study_site.id,             # StudySite-based scoping
        agreement_type=template.template_type,   # CDA / CTA / BUDGET / OTHER
        title=agreement_data.title,
        status=status_enum,
        created_by=current_user.get("user_id") if current_user else None,
        is_legacy=is_legacy,
    )
    logger.info(
        "Creating agreement %s with study_site_id=%s, study_id=%s, site_id=%s (site_id_external=%s)",
        str(agreement.id) if getattr(agreement, "id", None) else "<pending>",
        str(study_site.id),
        str(template.study_id),
        str(site.id),
        site.site_id,
    )
    db.add(agreement)
    await db.flush()
    
    # Clone template DOCX file and replace placeholders
    user_id = current_user.get("user_id") if current_user else None
    
    # Get template DOCX file path
    if not hasattr(template, 'template_file_path') or not template.template_file_path:
        raise HTTPException(
            status_code=400,
            detail="Template does not have a DOCX file. Please re-upload the template."
        )
    
    template_docx_path = Path(template.template_file_path)
    if not template_docx_path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"Template DOCX file not found: {template.template_file_path}"
        )
    
    # Create agreement documents directory
    agreement_dir = Path(settings.upload_dir) / "agreements" / str(agreement.id)
    agreement_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate document file path
    document_file_path = agreement_dir / f"version_1_{uuid.uuid4()}.docx"
    
    # Copy template and replace placeholders
    from app.utils.docx_placeholder_replace import replace_placeholders_in_docx
    
    sponsor_signatory_name = settings.sponsor_signatory_name
    sponsor_signatory_email = settings.sponsor_signatory_email
    current_user_email = current_user.get("email") if current_user else None
    
    # Get field_mappings from template (if available)
    field_mappings = template.field_mappings if hasattr(template, 'field_mappings') and template.field_mappings else None
    
    # Get placeholder_config for structural locking
    placeholder_config = template.placeholder_config if hasattr(template, 'placeholder_config') and template.placeholder_config else None
    
    replace_placeholders_in_docx(
        template_docx_path,
        profile,
        document_file_path,
        sponsor_signatory_name,
        sponsor_signatory_email,
        current_user_email,
        field_mappings=field_mappings,
        agreement=agreement,
        template_id=str(template.id),
        placeholder_config=placeholder_config
    )
    
    document = AgreementDocument(
        agreement_id=agreement.id,
        version_number=1,
        document_content=None,  # No longer storing JSON, using DOCX only
        document_file_path=str(document_file_path),  # Store DOCX file path
        document_html='',  # Empty string for DOCX-based documents (legacy field)
        created_from_template_id=template.id,
        created_by=user_id,
        is_signed_version='false',
    )
    db.add(document)
    await db.flush()
    
    # Set current_version_id to None for template-based agreements (we use documents instead)
    # Keep current_version_id for legacy compatibility
    
    # Create system comment for agreement creation with template
    await create_agreement_system_comment(
        db,
        agreement.id,
        None,
        f"Agreement draft created from template '{template.template_name}'"
    )
    
    await db.commit()
    
    # Reload with relationships
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement.id)
        .options(
            selectinload(Agreement.comments),
            selectinload(Agreement.documents),
            selectinload(Agreement.signed_documents),
        )
    )
    agreement = agreement_result.scalar_one()
    
    # Determine if user is internal
    user_id = current_user.get("user_id") if current_user else None
    is_internal = is_user_internal(user_id, db) if user_id else True
    
    return await build_agreement_response(db, agreement, is_internal, user_id)


@router.post("/agreements/{agreement_id}/versions", response_model=schemas.AgreementDocumentResponse)
async def upload_agreement_version(
    agreement_id: UUID,
    file: Optional[UploadFile] = File(None),
    document_html: Optional[str] = Form(None),
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Legacy endpoint preserved for backward compatibility.
    Runtime legacy file-based version uploads are no longer supported; callers should use DOCX-based flows.
    """
    raise HTTPException(
        status_code=400,
        detail="Legacy file-based versions are no longer supported. Use template-based agreements and document editor."
    )


@router.get("/agreements/{agreement_id}", response_model=schemas.AgreementResponse)
async def get_agreement_detail(
    agreement_id: UUID,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Get agreement detail including documents, comments, and permissions.
    """
    # Get agreement with all relationships
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(
            selectinload(Agreement.comments),
            selectinload(Agreement.documents).joinedload(AgreementDocument.inline_comments),
            selectinload(Agreement.signed_documents)
        )
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    # Determine if user is internal
    user_id = current_user.get("user_id") if current_user else None
    is_internal = is_user_internal(user_id, db) if user_id else True
    
    return await build_agreement_response(db, agreement, is_internal, user_id)


@router.get("/agreements/{agreement_id}/next-status")
async def get_next_allowed_status(
    agreement_id: UUID,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Get the next allowed status for an agreement.
    Returns null if no transitions are allowed (e.g., CLOSED).
    """
    agreement_result = await db.execute(
        select(Agreement).where(Agreement.id == agreement_id)
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    next_status = get_next_allowed_status_service(agreement.status)
    
    return {
        "current_status": agreement.status.value,
        "next_status": next_status.value if next_status else None,
        "can_transition": next_status is not None
    }


@router.patch("/agreements/{agreement_id}/status", response_model=schemas.AgreementResponse)
async def change_agreement_status(
    agreement_id: UUID,
    status_update: schemas.AgreementStatusUpdate,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Change agreement status using centralized workflow service.
    
    Enforces strict workflow rules:
    - Only allowed transitions are permitted
    - No backward transitions
    - No skipping states
    - Locking rules enforced
    
    Creates a system comment automatically with audit information.
    
    Enforces:
    - can_move_status permission check
    - Status transition validation
    """
    # Get agreement first to check permissions
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(
            selectinload(Agreement.documents),
            selectinload(Agreement.signed_documents)
        )
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    # Determine if user is internal for permission check
    user_id = current_user.get("user_id") if current_user else None
    is_internal = is_user_internal(user_id, db) if user_id else True
    
    # Check permissions
    permissions = get_agreement_permissions(agreement, is_internal)
    
    if not permissions["can_move_status"]:
        raise HTTPException(
            status_code=403,
            detail=f"Cannot change status. Agreement is locked or status '{agreement.status.value}' does not allow status transitions."
        )
    
    # Validate new status enum
    try:
        new_status_enum = AgreementStatusEnum(status_update.status.upper())
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid status: {status_update.status}")
    
    # Use centralized service function for status change
    try:
        agreement = await change_agreement_status_service(
            db,
            str(agreement_id),
            new_status_enum,
            user_id
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    
    # Reload with relationships for response
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement.id)
        .options(
            selectinload(Agreement.comments),
            selectinload(Agreement.documents),
            selectinload(Agreement.signed_documents)
        )
    )
    agreement = agreement_result.scalar_one()
    
    # Determine if user is internal
    user_id = current_user.get("user_id") if current_user else None
    is_internal = is_user_internal(user_id, db) if user_id else True
    
    return await build_agreement_response(db, agreement, is_internal, user_id)


# ---------------------------------------------------------------------------
# Agreement Workflow Reset Endpoint (for testing)
# ---------------------------------------------------------------------------

@router.delete("/agreements/{agreement_id}/reset")
async def reset_agreement_workflow(
    agreement_id: UUID,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Reset agreement workflow for testing purposes.
    This will completely reset the agreement:
    - Reset status to DRAFT
    - Delete all comments
    - Clear signature fields (zoho_request_id, signature_status)
    - Delete signed documents
    - Delete ALL document versions (AgreementDocument)
    - Delete ALL legacy file-based versions (historical only; table has been dropped in schema migration)
    - Reset current_version_id to None
    - This allows starting fresh with template selection
    """
    # Get agreement with all relationships
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(
            selectinload(Agreement.comments),
            selectinload(Agreement.documents),
            selectinload(Agreement.signed_documents),
        )
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    try:
        # Delete all comments
        await db.execute(
            delete(AgreementComment).where(AgreementComment.agreement_id == agreement_id)
        )
        logger.info(f"Deleted all comments for agreement {agreement_id}")
        
        # Delete signed documents
        await db.execute(
            delete(AgreementSignedDocument).where(AgreementSignedDocument.agreement_id == agreement_id)
        )
        logger.info(f"Deleted all signed documents for agreement {agreement_id}")
        
        # Delete ALL document versions (AgreementDocument)
        await db.execute(
            delete(AgreementDocument).where(AgreementDocument.agreement_id == agreement_id)
        )
        logger.info(f"Deleted all document versions for agreement {agreement_id}")
        
        # Also delete any DOCX/PDF files on disk for this agreement to avoid
        # stale/cached files being reused after reset.
        try:
            from pathlib import Path
            from app.config import settings as _settings
            import shutil

            agreement_dir = Path(_settings.upload_dir) / "agreements" / str(agreement_id)
            if agreement_dir.exists():
                shutil.rmtree(agreement_dir)
                logger.info(f"Deleted agreement files directory {agreement_dir} during reset for agreement {agreement_id}")
        except Exception as file_cleanup_error:
            logger.warning(
                "Failed to delete agreement files directory during reset for %s: %s",
                agreement_id,
                file_cleanup_error,
            )

        # Reset agreement fields
        agreement.status = AgreementStatusEnum.DRAFT
        agreement.zoho_request_id = None
        agreement.signature_status = None
        
        await db.commit()
        logger.info(f"Reset agreement {agreement_id} workflow to DRAFT - all versions deleted")
        
        # Reload agreement for response
        agreement_result = await db.execute(
            select(Agreement)
            .where(Agreement.id == agreement_id)
            .options(
                selectinload(Agreement.comments),
                selectinload(Agreement.documents),
                selectinload(Agreement.signed_documents)
            )
        )
        agreement = agreement_result.scalar_one()
        
        user_id = current_user.get("user_id") if current_user else None
        is_internal = is_user_internal(user_id, db) if user_id else True
        
        return {
            "status": "success",
            "message": "Agreement workflow reset successfully. All versions deleted. You can now select a template to start fresh.",
            "agreement": await build_agreement_response(db, agreement, is_internal, user_id)
        }
    
    except Exception as e:
        await db.rollback()
        logger.error(f"Failed to reset agreement workflow: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to reset workflow: {str(e)}")


@router.post("/agreements/{agreement_id}/comments", response_model=schemas.AgreementCommentResponse)
async def create_agreement_comment(
    agreement_id: UUID,
    comment_data: schemas.AgreementCommentCreate,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Create a new comment on an agreement.
    
    Permission rules:
    - INTERNAL users: Can create INTERNAL and EXTERNAL comments
    - EXTERNAL users: Can only create EXTERNAL comments
    - SYSTEM comments: Cannot be created manually
    
    Locking rules:
    - If agreement status = CLOSED: No new comments allowed
    
    Validations:
    - Comment type must be valid for user role
    - Content must not be empty (after trimming)
    - If version_id provided, must belong to same agreement
    """
    # Get agreement
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    # Locking rule: CLOSED agreements cannot have new comments
    if agreement.status == AgreementStatus.CLOSED:
        raise HTTPException(
            status_code=400,
            detail="Agreement is closed. No further comments allowed."
        )
    
    # Validate comment type enum
    try:
        comment_type_enum = CommentType(comment_data.comment_type.upper())
    except ValueError:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid comment type: {comment_data.comment_type}"
        )
    
    # Get user ID
    user_id = current_user.get("user_id") if current_user else None
    
    # Check permission to create this comment type
    can_create, error_msg = can_create_comment_type(user_id, comment_type_enum, db)
    if not can_create:
        raise HTTPException(status_code=403, detail=error_msg)
    
    # Validate and trim content
    content = comment_data.content.strip() if comment_data.content else ""
    if not content:
        raise HTTPException(status_code=400, detail="Comment content cannot be empty")
    
    # Create comment
    comment = AgreementComment(
        agreement_id=agreement.id,
        version_id=comment_data.version_id,
        comment_type=comment_type_enum,
        content=content,
        created_by=user_id  # Will be None for system comments, but we validate SYSTEM cannot be created manually
    )
    db.add(comment)
    await db.flush()
    await db.commit()
    await db.refresh(comment)
    
    return {
        "id": comment.id,
        "agreement_id": comment.agreement_id,
        "version_id": comment.version_id,
        "comment_type": comment.comment_type.value,
        "content": comment.content,
        "created_by": comment.created_by,
        "created_at": comment.created_at,
    }


@router.post("/agreements/{agreement_id}/create-from-template", response_model=schemas.AgreementDocumentResponse)
async def create_agreement_document_from_template(
    agreement_id: UUID,
    template_id: UUID = Form(...),
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Create a new document version from a template for an existing agreement.
    This is used when an agreement has been reset and needs a new document.
    """
    # Get agreement
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(selectinload(Agreement.documents))
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    # Validate template exists and is active
    template_result = await db.execute(
        select(StudyTemplate)
        .where(StudyTemplate.id == template_id)
        .where(StudyTemplate.is_active == 'true')
    )
    template = template_result.scalar_one_or_none()
    
    if not template:
        raise HTTPException(
            status_code=404,
            detail="Template not found or inactive"
        )
    
    # Check if agreement already has documents
    if agreement.documents and len(agreement.documents) > 0:
        raise HTTPException(
            status_code=400,
            detail="Agreement already has documents. Use save-document endpoint to create new versions."
        )
    
    # Check if agreement status allows document creation
    if agreement.status not in [AgreementStatus.DRAFT, AgreementStatus.UNDER_REVIEW]:
        raise HTTPException(
            status_code=403,
            detail=f"Cannot create document. Agreement status '{agreement.status.value}' does not allow document creation."
        )
    
    user_id = current_user.get("user_id") if current_user else None
    
    # Get template DOCX file path
    if not hasattr(template, 'template_file_path') or not template.template_file_path:
        raise HTTPException(
            status_code=400,
            detail="Template does not have a DOCX file. Please re-upload the template."
        )
    
    template_docx_path = Path(template.template_file_path)
    if not template_docx_path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"Template DOCX file not found: {template.template_file_path}"
        )
    
    # Fetch SiteProfile for placeholder merging
    profile_result = await db.execute(
        select(SiteProfile).where(SiteProfile.site_id == agreement.site_id)
    )
    profile = profile_result.scalar_one_or_none()
    
    if not profile:
        raise HTTPException(
            status_code=400,
            detail="SiteProfile not found. Please complete Site Profile before creating agreement."
        )
    
    # Create agreement documents directory
    agreement_dir = Path(settings.upload_dir) / "agreements" / str(agreement.id)
    agreement_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate document file path
    document_file_path = agreement_dir / f"version_1_{uuid.uuid4()}.docx"
    
    # Copy template and replace placeholders
    from app.utils.docx_placeholder_replace import replace_placeholders_in_docx
    
    sponsor_signatory_name = settings.sponsor_signatory_name
    sponsor_signatory_email = settings.sponsor_signatory_email
    current_user_email = current_user.get("email") if current_user else None
    
    # Get field_mappings from template (if available)
    field_mappings = template.field_mappings if hasattr(template, 'field_mappings') and template.field_mappings else None
    
    # Get placeholder_config for structural locking
    placeholder_config = template.placeholder_config if hasattr(template, 'placeholder_config') and template.placeholder_config else None
    
    replace_placeholders_in_docx(
        template_docx_path,
        profile,
        document_file_path,
        sponsor_signatory_name,
        sponsor_signatory_email,
        current_user_email,
        field_mappings=field_mappings,
        agreement=agreement,
        template_id=str(template.id),
        placeholder_config=placeholder_config
    )
    
    # Create first document version with DOCX file path
    document = AgreementDocument(
        agreement_id=agreement.id,
        version_number=1,
        document_content=None,  # No longer storing JSON, using DOCX only
        document_file_path=str(document_file_path),  # Store DOCX file path
        document_html='',  # Empty string for DOCX-based documents (legacy field)
        created_from_template_id=template.id,
        created_by=user_id,
        is_signed_version='false',
    )
    db.add(document)
    await db.flush()
    
    # Create system comment
    await create_agreement_system_comment(
        db,
        agreement.id,
        None,
        f"Agreement draft created from template '{template.template_name}'"
    )
    
    await db.commit()
    await db.refresh(document)
    
    return {
        "id": document.id,
        "agreement_id": document.agreement_id,
        "version_number": document.version_number,
        "document_content": document.document_content if document.document_content else {},
        "document_file_path": document.document_file_path if hasattr(document, 'document_file_path') else None,
        "created_from_template_id": str(document.created_from_template_id) if document.created_from_template_id else None,
        "created_by": document.created_by,
        "created_at": document.created_at,
        "is_signed_version": document.is_signed_version,
    }


@router.post("/agreements/{agreement_id}/save-document", response_model=schemas.AgreementDocumentResponse)
async def save_agreement_document(
    agreement_id: UUID,
    save_data: schemas.DocumentSaveRequest,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Save document content and create a new version.
    
    This endpoint:
    - Creates a new AgreementDocument version (does NOT overwrite previous)
    - Increments version_number
    - Creates SYSTEM comment: "Version X created after document edit"
    
    Validation rules:
    - Rejects if agreement.status in: READY_FOR_SIGNATURE, SENT_FOR_SIGNATURE, EXECUTED, CLOSED
    - Rejects if any document has is_signed_version = true
    
    Permissions:
    - Uses backend can_edit flag (no hardcoded status logic in frontend)
    """
    # Get agreement with documents
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(
            selectinload(Agreement.documents),
            selectinload(Agreement.signed_documents)
        )
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    # Determine if user is internal for permission check
    user_id = current_user.get("user_id") if current_user else None
    is_internal = is_user_internal(user_id, db) if user_id else True
    
    # Get permissions using comprehensive permission engine
    permissions = get_agreement_permissions(agreement, is_internal)
    
    # Enforce can_save permission (403 Forbidden for locked statuses)
    if not permissions["can_save"]:
        raise HTTPException(
            status_code=403,
            detail=f"Cannot save document. Agreement is locked or status '{agreement.status.value}' does not allow saving."
        )
    
    # Additional validation: Reject if signed version exists (double-check)
    has_signed_version = any(
        doc.is_signed_version == 'true' for doc in agreement.documents
    )
    if has_signed_version:
        raise HTTPException(
            status_code=403,
            detail="Cannot save document. A signed version exists."
        )
    
    # Get current max version number
    max_version_result = await db.execute(
        select(func.max(AgreementDocument.version_number))
        .where(AgreementDocument.agreement_id == agreement_id)
    )
    max_version_number = max_version_result.scalar_one_or_none()
    next_version_number = (max_version_number if max_version_number else 0) + 1
    
    # Get latest document to copy file path structure
    latest_document = max(agreement.documents, key=lambda d: d.version_number) if agreement.documents else None
    
    # Create agreement documents directory
    agreement_dir = Path(settings.upload_dir) / "agreements" / str(agreement.id)
    agreement_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate new document file path
    document_file_path = agreement_dir / f"version_{next_version_number}_{uuid.uuid4()}.docx"
    
    # For ONLYOFFICE, document is saved via callback, so we'll create the record
    # The actual file will be saved by the callback endpoint
    # For now, create placeholder (callback will update it)
    new_document = AgreementDocument(
        agreement_id=agreement.id,
        version_number=next_version_number,
        document_content=None,  # No longer storing JSON
        document_file_path=str(document_file_path),  # Will be updated by callback
        document_html='',  # Empty string for DOCX-based documents
        created_by=user_id,
        is_signed_version='false',
    )
    db.add(new_document)
    await db.flush()
    
    # Create SYSTEM comment
    await create_agreement_system_comment(
        db,
        agreement.id,
        None,
        f"Version {next_version_number} created after document edit"
    )
    
    await db.commit()
    await db.refresh(new_document)
    
    return {
        "id": new_document.id,
        "agreement_id": new_document.agreement_id,
        "version_number": new_document.version_number,
        "document_content": new_document.document_content if new_document.document_content else {},
        "document_file_path": new_document.document_file_path if hasattr(new_document, 'document_file_path') else None,
        "created_from_template_id": new_document.created_from_template_id,
        "created_by": new_document.created_by,
        "created_at": new_document.created_at,
        "is_signed_version": new_document.is_signed_version,
    }


@router.options("/agreements/{agreement_id}/onlyoffice-callback")
async def onlyoffice_callback_options(
    agreement_id: UUID,
):
    """Handle CORS preflight for ONLYOFFICE callback."""
    from fastapi.responses import Response
    return Response(
        headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
        }
    )

@router.post("/agreements/{agreement_id}/onlyoffice-callback")
async def onlyoffice_callback(
    agreement_id: UUID,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """
    ONLYOFFICE callback endpoint for document save.
    
    This endpoint receives the saved document from ONLYOFFICE when user saves.
    Creates a new version of the agreement document.
    """
    from app.utils.onlyoffice_utils import verify_jwt_token
    from fastapi.responses import JSONResponse
    import httpx
    
    try:
        # Get request body
        body = await request.json()

        # Log full payload (with token redacted) for debugging
        try:
            log_body = dict(body) if isinstance(body, dict) else {"raw": body}
            if "token" in log_body:
                log_body["token"] = "[REDACTED]"
        except Exception:
            log_body = {"raw": "unserializable body"}

        status = body.get("status", 0)
        logger.info(
            "ONLYOFFICE callback triggered for agreement %s with status=%s payload=%s",
            agreement_id,
            status,
            log_body,
        )

        # Verify JWT if enabled (ONLYOFFICE sends token in body, not header)
        if settings.onlyoffice_jwt_enabled and settings.onlyoffice_jwt_secret:
            token = body.get("token", "")
            if token:
                payload = verify_jwt_token(token)
                if not payload:
                    logger.warning("Invalid JWT token in ONLYOFFICE callback")
                    return JSONResponse(
                        content={"error": 1},
                        headers={
                            "Access-Control-Allow-Origin": "*",
                            "Access-Control-Allow-Methods": "POST, OPTIONS",
                            "Access-Control-Allow-Headers": "Content-Type, Authorization",
                        }
                    )
        
        # Get agreement and latest document
        agreement_result = await db.execute(
            select(Agreement)
            .where(Agreement.id == agreement_id)
            .options(selectinload(Agreement.documents))
        )
        agreement = agreement_result.scalar_one_or_none()
        
        if not agreement:
            raise HTTPException(status_code=404, detail="Agreement not found")
        
        # Get latest document
        if not agreement.documents:
            raise HTTPException(status_code=404, detail="No document found")
        
        latest_document = max(agreement.documents, key=lambda d: d.version_number)
        
        # Handle document save events
        # Status 1: Document is ready for editing (do NOT create version)
        # Status 2: Document is being saved (do NOT create version - just acknowledge)
        # Status 6: Document is being force saved (CREATE version only for this)
        # Status 7: Document save error (do NOT create version)
        
        # For status 2, just acknowledge - do NOT create version
        if status == 2:
            logger.info(
                "ONLYOFFICE callback status 2 (document saved) for agreement %s - acknowledging, NOT creating version",
                agreement_id
            )
            return JSONResponse(
                content={"error": 0},
                headers={
                    "Access-Control-Allow-Origin": "*",
                    "Access-Control-Allow-Methods": "POST, OPTIONS",
                    "Access-Control-Allow-Headers": "Content-Type, Authorization",
                }
            )
        
        # Only create new version for status 6 (manual force save)
        if status == 6:
            # Get document URL from callback
            url = body.get("url")
            if not url:
                logger.warning("ONLYOFFICE callback missing document URL")
                return JSONResponse(
                    content={"error": 0},
                    headers={
                        "Access-Control-Allow-Origin": "*",
                        "Access-Control-Allow-Methods": "POST, OPTIONS",
                        "Access-Control-Allow-Headers": "Content-Type, Authorization",
                    }
                )

            # Log the original URL from ONLYOFFICE
            logger.info("ONLYOFFICE callback document URL (original): %s", url)

            # Some ONLYOFFICE configs use 'localhost' in the URL, which is not
            # reachable from inside the backend container. Rewrite such hosts
            # to the internal ONLYOFFICE service host from settings.onlyoffice_url
            # (defaults to the docker-compose service 'onlyoffice:80').
            try:
                from urllib.parse import urlparse, urlunparse

                parsed = urlparse(url)
                target_host = parsed.hostname

                # Only ever rewrite localhost-style hosts. Public cloud hostnames
                # (e.g. *.azurewebsites.net) must be left untouched so DNS works.
                rewrite_hosts = {"localhost", "127.0.0.1"}

                if target_host in rewrite_hosts:
                    # Use the internal ONLYOFFICE URL (host:port) from settings.
                    internal_parsed = urlparse(settings.onlyoffice_url)
                    new_scheme = internal_parsed.scheme or parsed.scheme or "http"
                    new_netloc = internal_parsed.netloc or "onlyoffice"

                    rewritten = parsed._replace(scheme=new_scheme, netloc=new_netloc)
                    fixed_url = urlunparse(rewritten)
                    logger.info(
                        "Rewriting ONLYOFFICE document URL from %s to %s "
                        "for backend container connectivity.",
                        url,
                        fixed_url,
                    )
                    url = fixed_url
            except Exception as url_rewrite_error:
                logger.warning(
                    "Failed to analyze/rewrite ONLYOFFICE document URL %s: %s. "
                    "Proceeding with original URL.",
                    url,
                    str(url_rewrite_error),
                )

            # Download the saved document
            try:
                async with httpx.AsyncClient() as client:
                    response = await client.get(url, timeout=30.0)
                    response.raise_for_status()
                    docx_content = response.content
                logger.info(
                    "ONLYOFFICE document downloaded for agreement %s from %s (size=%s bytes, status=%s)",
                    agreement_id,
                    url,
                    len(docx_content),
                    response.status_code,
                )

                # PART 2: Duplicate save protection - compare file hashes
                # Compute hash of downloaded file
                new_file_hash = hashlib.sha256(docx_content).hexdigest()
                logger.info(
                    "Computed SHA256 hash of new document: %s (first 16 chars: %s)",
                    new_file_hash,
                    new_file_hash[:16]
                )

                # Compute hash of latest version file
                latest_file_hash = None
                if latest_document.document_file_path and Path(latest_document.document_file_path).exists():
                    try:
                        with open(latest_document.document_file_path, "rb") as f:
                            latest_file_content = f.read()
                        latest_file_hash = hashlib.sha256(latest_file_content).hexdigest()
                        logger.info(
                            "Computed SHA256 hash of latest version file: %s (first 16 chars: %s)",
                            latest_file_hash,
                            latest_file_hash[:16]
                        )
                    except Exception as hash_error:
                        logger.warning(
                            "Failed to compute hash of latest version file: %s. Proceeding with version creation.",
                            str(hash_error)
                        )

                # If hashes are identical, skip version creation
                if latest_file_hash and new_file_hash == latest_file_hash:
                    logger.info(
                        "No content change detected (hashes match). Skipping version creation for agreement %s",
                        agreement_id
                    )
                    return JSONResponse(
                        content={"error": 0},
                        headers={
                            "Access-Control-Allow-Origin": "*",
                            "Access-Control-Allow-Methods": "POST, OPTIONS",
                            "Access-Control-Allow-Headers": "Content-Type, Authorization",
                        }
                    )

                # Get next version number
                max_version_result = await db.execute(
                    select(func.max(AgreementDocument.version_number))
                    .where(AgreementDocument.agreement_id == agreement_id)
                )
                max_version_number = max_version_result.scalar_one_or_none()
                next_version_number = (max_version_number if max_version_number else 0) + 1

                logger.info(
                    "Creating version %s from manual save (status 6) for agreement %s",
                    next_version_number,
                    agreement_id
                )

                # Save new version
                agreement_dir = Path(settings.upload_dir) / "agreements" / str(agreement.id)
                agreement_dir.mkdir(parents=True, exist_ok=True)
                document_file_path = agreement_dir / f"version_{next_version_number}_{uuid.uuid4()}.docx"

                with open(document_file_path, "wb") as f:
                    f.write(docx_content)
                
                # Create new document version
                new_document = AgreementDocument(
                    agreement_id=agreement.id,
                    version_number=next_version_number,
                    document_content=None,
                    document_file_path=str(document_file_path),
                    document_html='',
                    created_by=latest_document.created_by,
                    is_signed_version='false',
                )
                db.add(new_document)
                
                # Create system comment
                await create_agreement_system_comment(
                    db,
                    agreement.id,
                    None,
                    f"Version {next_version_number} created after document edit via ONLYOFFICE"
                )
                
                await db.commit()
                logger.info(
                    "Saved document version %s for agreement %s (hash: %s)",
                    next_version_number,
                    agreement_id,
                    new_file_hash[:16]
                )
            
            except Exception as e:
                logger.error(
                    "Failed to save document from ONLYOFFICE callback for agreement %s: %s",
                    agreement_id,
                    str(e),
                    exc_info=True,
                )
                await db.rollback()
                return JSONResponse(
                    content={"error": 1},
                    headers={
                        "Access-Control-Allow-Origin": "*",
                        "Access-Control-Allow-Methods": "POST, OPTIONS",
                        "Access-Control-Allow-Headers": "Content-Type, Authorization",
                    }
                )
        
        # Return success with CORS headers
        return JSONResponse(
            content={"error": 0},
            headers={
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "POST, OPTIONS",
                "Access-Control-Allow-Headers": "Content-Type, Authorization",
            }
        )
        
    except Exception as e:
        logger.error(f"ONLYOFFICE callback error: {str(e)}")
        return JSONResponse(
            content={"error": 1},
            headers={
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "POST, OPTIONS",
                "Access-Control-Allow-Headers": "Content-Type, Authorization",
            }
        )


@router.get("/agreements/{agreement_id}/onlyoffice-config")
async def get_onlyoffice_config(
    agreement_id: UUID,
    request: Request,
    version: Optional[int] = None,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Get ONLYOFFICE editor configuration for an agreement document.
    
    Returns:
        Configuration object for ONLYOFFICE editor initialization
    """
    from app.utils.onlyoffice_utils import create_document_config, get_onlyoffice_editor_url
    from fastapi.responses import JSONResponse
    
    # Get agreement and documents
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(selectinload(Agreement.documents))
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    if not agreement.documents:
        raise HTTPException(status_code=404, detail="No document found")

    # Select document based on requested version, or latest if not specified
    if version is not None:
        document = next(
            (d for d in agreement.documents if d.version_number == version),
            None,
        )
        if not document:
            raise HTTPException(
                status_code=404,
                detail=f"Document version {version} not found",
            )
    else:
        document = max(agreement.documents, key=lambda d: d.version_number)
    
    # Check if document has DOCX file
    if not hasattr(document, 'document_file_path') or not document.document_file_path:
        raise HTTPException(
            status_code=400,
            detail="Document does not have a DOCX file. Please create document from template."
        )
    
    # Determine edit mode based on agreement status
    # Locking rules:
    # - SENT_FOR_SIGNATURE: Read-only (view mode)
    # - EXECUTED: Disable editor entirely (return error)
    if agreement.status == AgreementStatusEnum.EXECUTED:
        raise HTTPException(
            status_code=403,
            detail="Cannot open editor. Agreement is EXECUTED and cannot be edited."
        )
    
    permissions = get_agreement_permissions(agreement, True)  # Assume internal for now
    # If status is SENT_FOR_SIGNATURE, force view mode
    if agreement.status == AgreementStatusEnum.SENT_FOR_SIGNATURE:
        mode = "view"
    else:
        mode = "view" if not permissions["can_edit"] else "edit"
    
    # Get user info
    user_id = current_user.get("user_id", "anonymous") if current_user else "anonymous"
    user_name = current_user.get("name", "User") if current_user else "User"
    
    # Create document and callback URLs
    # IMPORTANT (ONLYOFFICE behavior):
    # - Document URL is fetched by ONLYOFFICE Document Server (in the Docker container),
    #   NOT directly by the browser, so it must be reachable from the container.
    # - Callback URL is also called from the Document Server.
    # - Backend and ONLYOFFICE may be in the same docker-compose network (local dev),
    #   or reachable via a public URL in production (Azure App Service).
    # - Use a configurable backend_internal_url so that:
    #     * Local default is http://backend:8000 (docker-compose service name)
    #     * Production can set BACKEND_INTERNAL_URL to the public HTTPS URL.
    
    # Use configured backend internal URL for ONLYOFFICE callbacks
    internal_base = settings.backend_internal_url
    
    # Both document and callback URLs must be reachable from ONLYOFFICE container
    document_url = f"{internal_base}/api/agreements/{agreement_id}/document-file"
    if version is not None:
        document_url = f"{document_url}?version={version}"
    callback_url = f"{internal_base}/api/agreements/{agreement_id}/onlyoffice-callback"
    
    logger.info(f"ONLYOFFICE config - document_url: {document_url}, callback_url: {callback_url}")
    
    # Generate document key (unique identifier for ONLYOFFICE)
    # Use stable document.id to prevent editor remounting on every state change
    # This ensures the editor stays mounted and doesn't glitch
    document_key = str(document.id)
    
    # Create configuration
    config = create_document_config(
        document_url=document_url,
        callback_url=callback_url,
        document_key=document_key,
        document_title=agreement.title,
        user_id=user_id,
        user_name=user_name,
        mode=mode,
    )
    
    editor_url = get_onlyoffice_editor_url()
    logger.info(
        "ONLYOFFICE config - document_url: %s, callback_url: %s, editor_url: %s, version=%s",
        document_url,
        callback_url,
        editor_url,
        version if version is not None else document.version_number,
    )

    return JSONResponse(
        content={
            "editorUrl": editor_url,
            "config": config,
        }
    )


@router.get("/agreements/{agreement_id}/document-file")
async def get_agreement_document_file(
    agreement_id: UUID,
    version: Optional[int] = None,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Get agreement document DOCX file for ONLYOFFICE.
    
    Returns:
        DOCX file response
    """
    from fastapi.responses import FileResponse
    
    # Get agreement and documents
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(selectinload(Agreement.documents))
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    if not agreement.documents:
        raise HTTPException(status_code=404, detail="No document found")
    
    # Get document (latest or specified version)
    if version:
        document = next((d for d in agreement.documents if d.version_number == version), None)
    else:
        document = max(agreement.documents, key=lambda d: d.version_number)
    
    if not document:
        raise HTTPException(status_code=404, detail="Document version not found")
    
    # Check if document has DOCX file
    if not hasattr(document, 'document_file_path') or not document.document_file_path:
        raise HTTPException(
            status_code=404,
            detail="Document file not found"
        )
    
    file_path = Path(document.document_file_path)
    if not file_path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"Document file not found at: {document.document_file_path}"
        )
    
    # Add CORS headers for ONLYOFFICE to access the document
    headers = {
        "Access-Control-Allow-Origin": "*",  # ONLYOFFICE needs to access this
        "Access-Control-Allow-Methods": "GET, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type, Authorization",
    }
    
    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"agreement_{agreement_id}_v{document.version_number}.docx",
        headers=headers
    )


@router.post("/agreements/{agreement_id}/send-for-signature")
async def send_agreement_for_signature(
    agreement_id: UUID,
    site_signer_email: Optional[str] = Form(None, description="Site signatory email (auto-filled from SiteProfile, signs first)"),
    sponsor_signer_email: Optional[str] = Form(None, description="Sponsor/Internal signatory email (auto-filled from config, signs second)"),
    cc_emails: Optional[str] = Form(None, description="CC recipients (comma-separated)"),
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Send agreement for signature via Zoho Sign.
    
    Requirements:
    - agreement.status = READY_FOR_SIGNATURE
    - No signed version exists
    - Latest AgreementDocument exists
    
    Process:
    1. Generate PDF from latest AgreementDocument (TipTap JSON -> HTML -> PDF)
    2. Send to Zoho Sign with signing order (site first, sponsor second)
    3. Store zoho_request_id and signature_status
    4. Update agreement.status to SENT_FOR_SIGNATURE
    5. Create SYSTEM comment
    """
    from app.services.zoho_sign_service import zoho_sign_service
    import tempfile
    
    # Get agreement with documents
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(
            selectinload(Agreement.documents),
            selectinload(Agreement.signed_documents)
        )
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    # Validate status
    if agreement.status != AgreementStatusEnum.READY_FOR_SIGNATURE:
        raise HTTPException(
            status_code=400,
            detail=f"Agreement must be in READY_FOR_SIGNATURE status. Current status: {agreement.status.value}"
        )
    
    # Check if signed version exists
    has_signed_version = any(
        doc.is_signed_version == 'true' for doc in agreement.documents
    )
    if has_signed_version:
        raise HTTPException(
            status_code=400,
            detail="Cannot send for signature: A signed version already exists."
        )
    
    # Check if zoho_request_id already exists (already sent)
    if agreement.zoho_request_id:
        raise HTTPException(
            status_code=400,
            detail=f"Agreement already sent for signature. Request ID: {agreement.zoho_request_id}"
        )
    
    # Get latest document
    if not agreement.documents or len(agreement.documents) == 0:
        raise HTTPException(
            status_code=400,
            detail="No document versions found. Cannot send for signature."
        )
    
    latest_document = max(agreement.documents, key=lambda d: d.version_number)
    
    # Auto-fill recipients from SiteProfile and config
    from app.models import SiteProfile
    
    # Get site profile for auto-filling site signatory
    site_result = await db.execute(
        select(Site).where(Site.id == agreement.site_id)
    )
    site = site_result.scalar_one_or_none()
    
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")
    
    profile_result = await db.execute(
        select(SiteProfile).where(SiteProfile.site_id == site.id)
    )
    profile = profile_result.scalar_one_or_none()
    
    if not profile:
        raise HTTPException(
            status_code=400,
            detail="SiteProfile not found. Cannot auto-fill site signatory email."
        )
    
    # Auto-fill site signatory email from SiteProfile
    if not site_signer_email:
        site_signer_email = profile.authorized_signatory_email
        if not site_signer_email:
            raise HTTPException(
                status_code=400,
                detail="Site signatory email not found in SiteProfile. Please provide site_signer_email."
            )
    
    # Auto-fill sponsor signatory email from config or current user
    if not sponsor_signer_email:
        sponsor_signer_email = settings.sponsor_signatory_email
        if not sponsor_signer_email:
            # Fallback to current user email
            if current_user:
                sponsor_signer_email = current_user.get("email")
            if not sponsor_signer_email:
                raise HTTPException(
                    status_code=400,
                    detail="Sponsor signatory email not configured. Please set SPONSOR_SIGNATORY_EMAIL in environment or provide sponsor_signer_email."
                )
    
    # Validate emails
    if not site_signer_email or not site_signer_email.strip():
        raise HTTPException(status_code=400, detail="Site signatory email is required")
    if not sponsor_signer_email or not sponsor_signer_email.strip():
        raise HTTPException(status_code=400, detail="Sponsor signatory email is required")
    
    # Check if document has DOCX file
    if not hasattr(latest_document, 'document_file_path') or not latest_document.document_file_path:
        raise HTTPException(
            status_code=400,
            detail="Document does not have a DOCX file. Cannot generate PDF."
        )
    
    docx_path = Path(latest_document.document_file_path)
    if not docx_path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"Document file not found: {latest_document.document_file_path}"
        )
    
    # Convert DOCX to PDF
    try:
        from app.utils.docx_to_pdf import docx_to_pdf
        
        # Create temporary file for PDF
        temp_dir = Path(settings.upload_dir) / "temp_agreements"
        temp_dir.mkdir(parents=True, exist_ok=True)
        temp_pdf_path = temp_dir / f"agreement_{agreement.id}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.pdf"
        
        docx_to_pdf(docx_path, temp_pdf_path, method="libreoffice")
        logger.info(f"Generated PDF from DOCX for agreement {agreement_id}: {temp_pdf_path}")
    except Exception as e:
        logger.error(f"Failed to generate PDF: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate PDF: {str(e)}"
        )
    
    # Prepare recipients for Zoho Sign
    recipients = [
        {
            "name": "Site Signatory",
            "email": site_signer_email.strip(),
            "action": "SIGN",
            "signing_order": 1,
        },
        {
            "name": "Sponsor Signatory",
            "email": sponsor_signer_email.strip(),
            "action": "SIGN",
            "signing_order": 2,
        },
    ]
    
    # Prepare CC recipients
    cc_recipients = None
    if cc_emails:
        cc_recipients = [email.strip() for email in cc_emails.split(",") if email.strip()]
    
    # Send to Zoho Sign
    try:
        request_name = f"Agreement: {agreement.title}"
        message = f"Please review and sign the agreement: {agreement.title}"
        
        zoho_response = zoho_sign_service.create_signature_request(
            request_name=request_name,
            document_path=temp_pdf_path,
            recipients=recipients,
            cc_recipients=cc_recipients,
            message=message,
        )
        
        zoho_request_id = zoho_response.get("requests", {}).get("request_id")
        if not zoho_request_id:
            raise Exception("Zoho Sign response missing request_id")
        
        logger.info(f"Sent agreement {agreement_id} to Zoho Sign. Request ID: {zoho_request_id}")
        
    except Exception as e:
        logger.error(f"Failed to send to Zoho Sign: {str(e)}")
        # Clean up temp PDF
        if temp_pdf_path.exists():
            try:
                temp_pdf_path.unlink()
            except:
                pass
        raise HTTPException(
            status_code=500,
            detail=f"Failed to send to Zoho Sign: {str(e)}"
        )
    
    # Update agreement
    agreement.zoho_request_id = zoho_request_id
    agreement.signature_status = "SENT"
    agreement.status = AgreementStatusEnum.SENT_FOR_SIGNATURE
    
    # Create SYSTEM comment
    await create_agreement_system_comment(
        db,
        agreement.id,
        None,
        f"Agreement sent for signature via Zoho Sign. Request ID: {zoho_request_id}"
    )
    
    await db.commit()
    
    # Create Notice Board entry for SENT_FOR_SIGNATURE status
    try:
        user_id = current_user.get("user_id") if current_user else None
        await create_agreement_notice_board_entry(
            db,
            agreement,
            AgreementStatusEnum.SENT_FOR_SIGNATURE.value,
            created_by=user_id
        )
    except Exception as e:
        logger.error(f"Failed to create Notice Board entry for agreement {agreement_id}: {str(e)}", exc_info=True)
        # Don't fail the transaction if notice creation fails
    
    # Clean up temp PDF (Zoho has it now)
    if temp_pdf_path.exists():
        try:
            temp_pdf_path.unlink()
        except:
            pass
    
    # Reload agreement for response
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(
            selectinload(Agreement.comments),
            selectinload(Agreement.documents),
            selectinload(Agreement.signed_documents)
        )
    )
    agreement = agreement_result.scalar_one()
    
    user_id = current_user.get("user_id") if current_user else None
    is_internal = is_user_internal(user_id, db) if user_id else True
    
    return await build_agreement_response(db, agreement, is_internal, user_id)


@router.post("/agreements/{agreement_id}/sync-zoho-status")
async def sync_agreement_zoho_status(
    agreement_id: UUID,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Manually sync Zoho Sign status for an agreement.
    This endpoint fetches the current status from Zoho Sign and updates the CRM.
    Useful when webhooks are not configured or not working.
    """
    from app.services.zoho_sign_service import zoho_sign_service
    from pathlib import Path
    
    # Get agreement
    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(
            selectinload(Agreement.documents),
            selectinload(Agreement.signed_documents)
        )
    )
    agreement = agreement_result.scalar_one_or_none()
    
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    if not agreement.zoho_request_id:
        raise HTTPException(status_code=400, detail="No Zoho Sign request ID found for this agreement")
    
    logger.info(f"Syncing Zoho Sign status for agreement {agreement_id}, request_id: {agreement.zoho_request_id}")
    
    try:
        # Fetch current status from Zoho Sign
        request_details = zoho_sign_service.get_request_details(agreement.zoho_request_id)
        requests_data = request_details.get("requests", {})
        
        # Zoho's response structure can vary slightly; normalize it
        if isinstance(requests_data, list) and requests_data:
            requests_data = requests_data[0]
        
        raw_status = (
            requests_data.get("status")
            or requests_data.get("request_status")
            or requests_data.get("request_status_name")
            or requests_data.get("request_status_text")
        )
        zoho_status = (str(raw_status) if raw_status is not None else "").upper()
        
        logger.info(
            f"Zoho Sign status for {agreement.zoho_request_id}: raw={raw_status!r}, normalized={zoho_status}"
        )
        
        # Update based on Zoho status
        if zoho_status == "COMPLETED":
            # Download signed PDF and process (same logic as webhook)
            signed_dir = Path(settings.upload_dir) / "signed_agreements"
            signed_dir.mkdir(parents=True, exist_ok=True)
            signed_pdf_path = signed_dir / f"agreement_{agreement.id}_signed_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.pdf"
            
            try:
                zoho_sign_service.download_signed_document(agreement.zoho_request_id, signed_pdf_path)
                logger.info(f"Downloaded signed agreement PDF: {signed_pdf_path}")
                
                # Get signed_at from Zoho response (if available)
                signed_at = None
                if requests_data.get("completed_time"):
                    try:
                        signed_at = datetime.fromisoformat(requests_data["completed_time"].replace("Z", "+00:00"))
                    except:
                        pass
                
                # Create AgreementSignedDocument record
                signed_document = AgreementSignedDocument(
                    agreement_id=agreement.id,
                    file_path=str(signed_pdf_path),
                    signed_at=signed_at,
                    zoho_request_id=agreement.zoho_request_id,
                )
                db.add(signed_document)
                await db.flush()
                
                # Get latest document version number
                max_version_result = await db.execute(
                    select(func.max(AgreementDocument.version_number))
                    .where(AgreementDocument.agreement_id == agreement.id)
                )
                max_version_number = max_version_result.scalar_one_or_none()
                next_version_number = (max_version_number if max_version_number else 0) + 1
                
                # Get latest document content
                latest_document = max(agreement.documents, key=lambda d: d.version_number)
                
                # Create new AgreementDocument version marked as signed
                signed_document_version = AgreementDocument(
                    agreement_id=agreement.id,
                    version_number=next_version_number,
                    document_content=latest_document.document_content,
                    document_html='',
                    created_from_template_id=None,
                    created_by=None,
                    is_signed_version='true',
                )
                db.add(signed_document_version)
                await db.flush()
                
                # Update agreement status
                agreement.status = AgreementStatusEnum.EXECUTED
                agreement.signature_status = "COMPLETED"
                
                # Create SYSTEM comment
                # Note: version_id is for legacy file-based versions only, not AgreementDocument
                # For document-based agreements, we pass None
                await create_agreement_system_comment(
                    db,
                    agreement.id,
                    None,  # version_id is only for legacy file-based versions, not AgreementDocument
                    "Agreement executed. Signed document received from Zoho Sign (synced)."
                )
                
                # Create Notice Board entry for EXECUTED status
                try:
                    await create_agreement_notice_board_entry(
                        db,
                        agreement,
                        AgreementStatusEnum.EXECUTED.value,
                        created_by=user_id if current_user else None
                    )
                except Exception as e:
                    logger.error(f"Failed to create Notice Board entry for agreement {agreement_id}: {str(e)}", exc_info=True)
                    # Don't fail the transaction if notice creation fails
                
                # PART 2: Update Site Status Milestone - Complete CDA Execution
                # Only for CDA agreements
                try:
                    await complete_cda_execution_milestone(
                        db,
                        agreement,
                        created_by=user_id if current_user else None
                    )
                except Exception as e:
                    logger.error(f"Failed to complete CDA Execution milestone for agreement {agreement_id}: {str(e)}", exc_info=True)
                    # Don't fail the transaction if milestone update fails
                
                await db.commit()
                logger.info(f"Agreement {agreement_id} marked as EXECUTED after sync")
                
            except Exception as e:
                logger.error(f"Failed to download signed document during sync: {str(e)}")
                # Still update status even if download fails
                agreement.signature_status = "COMPLETED"
                agreement.status = AgreementStatusEnum.EXECUTED
                await create_agreement_system_comment(
                    db,
                    agreement.id,
                    None,
                    f"Agreement signature completed (synced), but PDF download failed: {str(e)}"
                )
                await db.commit()
        
        elif zoho_status == "DECLINED":
            agreement.signature_status = "DECLINED"
            await create_agreement_system_comment(
                db,
                agreement.id,
                None,
                "Agreement signature request was declined (synced from Zoho)."
            )
            await db.commit()
            logger.info(f"Agreement {agreement_id} signature declined (synced)")
        
        elif zoho_status == "EXPIRED":
            agreement.signature_status = "EXPIRED"
            await create_agreement_system_comment(
                db,
                agreement.id,
                None,
                "Agreement signature request expired (synced from Zoho)."
            )
            await db.commit()
            logger.info(f"Agreement {agreement_id} signature expired (synced)")
        
        # Reload agreement for response
        agreement_result = await db.execute(
            select(Agreement)
            .where(Agreement.id == agreement_id)
            .options(
                selectinload(Agreement.comments),
                selectinload(Agreement.documents),
                selectinload(Agreement.signed_documents)
            )
        )
        agreement = agreement_result.scalar_one()
        
        user_id = current_user.get("user_id") if current_user else None
        is_internal = is_user_internal(user_id, db) if user_id else True
        
        # Build full agreement response with signed documents
        agreement_response = await build_agreement_response(db, agreement, is_internal, user_id)
        
        return {
            "status": "success",
            "message": f"Synced Zoho Sign status: {zoho_status}",
            "zoho_status": zoho_status,
            "agreement_status": agreement.status.value,
            "agreement": agreement_response
        }
    
    except Exception as e:
        logger.error(f"Failed to sync Zoho Sign status: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to sync status: {str(e)}")


@router.get("/agreements/{agreement_id}/signed-document/{signed_document_id}")
async def download_signed_agreement_pdf(
    agreement_id: UUID,
    signed_document_id: UUID,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Download signed agreement PDF.
    
    Args:
        agreement_id: Agreement UUID
        signed_document_id: AgreementSignedDocument UUID
        
    Returns:
        PDF file response
    """
    # Get signed document
    signed_doc_result = await db.execute(
        select(AgreementSignedDocument)
        .where(AgreementSignedDocument.id == signed_document_id)
        .where(AgreementSignedDocument.agreement_id == agreement_id)
    )
    signed_doc = signed_doc_result.scalar_one_or_none()
    
    if not signed_doc:
        raise HTTPException(status_code=404, detail="Signed document not found")
    
    # Check if file exists
    pdf_path = Path(signed_doc.file_path)
    if not pdf_path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"Signed PDF file not found at: {signed_doc.file_path}"
        )
    
    # Return file
    filename = pdf_path.name
    return FileResponse(
        path=str(pdf_path),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'inline; filename="{filename}"'
        },
    )


# ---------------------------------------------------------------------------
# Study Template Management Endpoints
# ---------------------------------------------------------------------------

@router.post("/studies/{study_id}/templates", response_model=schemas.StudyTemplateResponse)
async def create_study_template(
    study_id: UUID,
    template_name: str = Form(...),
    template_type: str = Form(...),
    template_file: UploadFile = File(...),
    field_mappings: Optional[str] = Form(None),  # JSON string for field mappings
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Create a new study template from DOCX upload.
    DOCX files are automatically converted to HTML and then to TipTap JSON format.
    Preserves formatting, tables, images, and layout.
    """
    # Validate study exists
    study_result = await db.execute(
        select(Study).where(Study.id == study_id)
    )
    study = study_result.scalar_one_or_none()
    
    if not study:
        raise HTTPException(status_code=404, detail="Study not found")
    
    # Validate template type enum
    try:
        template_type_enum = TemplateType(template_type.upper())
    except ValueError:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid template type: {template_type}"
        )
    
    # Validate file type - DOCX only
    if not template_file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_extension = Path(template_file.filename).suffix.lower()
    if file_extension not in ['.docx', '.doc']:
        raise HTTPException(
            status_code=400,
            detail="Only DOCX files are supported. Please upload a .docx file. PDF uploads are no longer supported."
        )
    
    # Save uploaded DOCX file permanently
    upload_dir = Path(settings.upload_dir) / "templates"
    upload_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate unique filename
    file_id = uuid.uuid4()
    file_name = f"template_{file_id}.docx"
    template_file_path = upload_dir / file_name
    
    try:
        # Save uploaded file
        with open(template_file_path, "wb") as buffer:
            shutil.copyfileobj(template_file.file, buffer)
        
        logger.info(f"Saved template DOCX file: {template_file_path}")
        
        # For backward compatibility, still detect placeholders from DOCX content
        # We'll read the DOCX to detect placeholders for configuration
        from docx import Document
        doc = Document(template_file_path)
        placeholder_text = ""
        for paragraph in doc.paragraphs:
            placeholder_text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    placeholder_text += cell.text + "\n"
        
        # Detect placeholders using regex
        import re
        placeholder_pattern = r'\{\{([A-Z_]+)\}\}'
        detected_placeholders = set(re.findall(placeholder_pattern, placeholder_text))
        
        # Create placeholder configuration
        placeholder_config = {}
        for p_name in sorted(list(detected_placeholders)):
            placeholder_config[p_name] = {"editable": True}
        
        # Parse field_mappings if provided
        parsed_field_mappings = None
        if field_mappings:
            try:
                import json
                parsed_field_mappings = json.loads(field_mappings)
                if not isinstance(parsed_field_mappings, dict):
                    raise ValueError("field_mappings must be a JSON object")
                logger.info(f"Parsed field_mappings: {parsed_field_mappings}")
            except json.JSONDecodeError as e:
                logger.warning(f"Invalid JSON in field_mappings: {e}")
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid JSON in field_mappings: {str(e)}"
                )
            except Exception as e:
                logger.warning(f"Error parsing field_mappings: {e}")
                raise HTTPException(
                    status_code=400,
                    detail=f"Error parsing field_mappings: {str(e)}"
                )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to process DOCX template: {str(e)}")
        # Clean up file if processing failed
        if template_file_path.exists():
            template_file_path.unlink()
        raise HTTPException(
            status_code=400,
            detail=f"Failed to process DOCX file: {str(e)}"
        )
    
    # Create template with DOCX file path
    template = StudyTemplate(
        study_id=study.id,
        template_name=template_name,
        template_type=template_type_enum,
        template_content=None,  # No longer storing JSON, using DOCX only
        template_file_path=str(template_file_path),  # Store DOCX file path
        document_html='',  # Empty string for DOCX-based templates (legacy field)
        placeholder_config=placeholder_config,  # Auto-detected placeholder configuration
        field_mappings=parsed_field_mappings,  # Dynamic field mappings
        created_by=current_user.get("user_id") if current_user else None,
        is_active='true',
    )
    db.add(template)
    await db.flush()
    await db.commit()
    await db.refresh(template)
    
    return {
        "id": template.id,
        "study_id": template.study_id,
        "template_name": template.template_name,
        "template_type": template.template_type.value,
        "template_content": template.template_content if hasattr(template, 'template_content') and template.template_content else {"type": "doc", "content": []},
        "placeholder_config": template.placeholder_config if hasattr(template, 'placeholder_config') and template.placeholder_config else {},
        "created_by": template.created_by,
        "created_at": template.created_at,
        "updated_at": template.updated_at,
        "is_active": template.is_active,
    }


@router.get("/studies/{study_id}/templates", response_model=List[schemas.StudyTemplateResponse])
async def list_study_templates(
    study_id: UUID,
    active_only: bool = Query(True, description="Return only active templates"),
    db: AsyncSession = Depends(get_db),
):
    """
    List all templates for a study.
    """
    # Validate study exists
    study_result = await db.execute(
        select(Study).where(Study.id == study_id)
    )
    study = study_result.scalar_one_or_none()
    
    if not study:
        raise HTTPException(status_code=404, detail="Study not found")
    
    # Query templates
    query = select(StudyTemplate).where(StudyTemplate.study_id == study_id)
    if active_only:
        query = query.where(StudyTemplate.is_active == 'true')
    query = query.order_by(StudyTemplate.created_at.desc())
    
    templates_result = await db.execute(query)
    templates = templates_result.scalars().all()
    
    return [
        {
            "id": template.id,
            "study_id": template.study_id,
            "template_name": template.template_name,
            "template_type": template.template_type.value,
            "template_content": template.template_content if hasattr(template, 'template_content') and template.template_content else {"type": "doc", "content": []},
            "template_file_path": template.template_file_path if hasattr(template, 'template_file_path') else None,
            "placeholder_config": template.placeholder_config if hasattr(template, 'placeholder_config') and template.placeholder_config else {},
            "field_mappings": template.field_mappings if hasattr(template, 'field_mappings') and template.field_mappings else None,
            "created_by": template.created_by,
            "created_at": template.created_at,
            "updated_at": template.updated_at,
            "is_active": template.is_active,
        }
        for template in templates
    ]


@router.get("/templates/field-mapping-options")
async def get_template_field_mapping_options():
    """
    Get available data source fields for dynamic field mappings.

    Returns a static list of Site Profile and Agreement fields that can be
    mapped to placeholders in templates.
    """
    return {
        "site_profile": [
            {"field": "site_name", "label": "Site Name"},
            {"field": "hospital_name", "label": "Hospital Name"},
            {"field": "pi_name", "label": "PI Name"},
            {"field": "pi_email", "label": "PI Email"},
            {"field": "pi_phone", "label": "PI Phone"},
            {"field": "primary_contracting_entity", "label": "Primary Contracting Entity"},
            {"field": "authorized_signatory_name", "label": "Authorized Signatory Name"},
            {"field": "authorized_signatory_email", "label": "Authorized Signatory Email"},
            {"field": "authorized_signatory_title", "label": "Authorized Signatory Title"},
            {"field": "address_line_1", "label": "Address Line 1"},
            {"field": "city", "label": "City"},
            {"field": "state", "label": "State"},
            {"field": "country", "label": "Country"},
            {"field": "postal_code", "label": "Postal Code"},
            {"field": "site_coordinator_name", "label": "Site Coordinator Name"},
            {"field": "site_coordinator_email", "label": "Site Coordinator Email"},
        ],
        "agreement": [
            {"field": "title", "label": "Agreement Title"},
            {"field": "status", "label": "Agreement Status"},
            {"field": "created_by", "label": "Created By"},
        ],
    }


@router.get("/templates/{template_id}", response_model=schemas.StudyTemplateResponse)
async def get_template_detail(
    template_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    """
    Get template detail by ID.
    """
    template_result = await db.execute(
        select(StudyTemplate).where(StudyTemplate.id == template_id)
    )
    template = template_result.scalar_one_or_none()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    return {
        "id": template.id,
        "study_id": template.study_id,
        "template_name": template.template_name,
        "template_type": template.template_type.value,
        "template_content": template.template_content if hasattr(template, 'template_content') and template.template_content else {"type": "doc", "content": []},
        "template_file_path": template.template_file_path if hasattr(template, 'template_file_path') else None,
        "placeholder_config": template.placeholder_config if hasattr(template, 'placeholder_config') and template.placeholder_config else {},
        "field_mappings": template.field_mappings if hasattr(template, 'field_mappings') and template.field_mappings else None,
        "created_by": template.created_by,
        "created_at": template.created_at,
        "updated_at": template.updated_at,
        "is_active": template.is_active,
    }


@router.get("/templates/{template_id}/template-file")
async def get_template_file(
    template_id: UUID,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Get template DOCX file for ONLYOFFICE viewing.
    
    Returns:
        DOCX file response
    """
    from fastapi.responses import FileResponse
    
    template_result = await db.execute(
        select(StudyTemplate).where(StudyTemplate.id == template_id)
    )
    template = template_result.scalar_one_or_none()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Check if template has DOCX file
    if not hasattr(template, 'template_file_path') or not template.template_file_path:
        raise HTTPException(
            status_code=404,
            detail="Template file not found"
        )
    
    file_path = Path(template.template_file_path)
    if not file_path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"Template file not found at: {template.template_file_path}"
        )
    
    # Add CORS headers for ONLYOFFICE to access the document
    headers = {
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Methods": "GET, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type, Authorization",
    }
    
    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"template_{template_id}_{template.template_name}.docx",
        headers=headers
    )


@router.get("/templates/{template_id}/onlyoffice-config")
async def get_template_onlyoffice_config(
    template_id: UUID,
    request: Request,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Get ONLYOFFICE editor configuration for viewing a template document.
    
    Returns:
        Configuration object for ONLYOFFICE editor initialization (read-only mode)
    """
    from app.utils.onlyoffice_utils import create_document_config, get_onlyoffice_editor_url
    from fastapi.responses import JSONResponse
    
    template_result = await db.execute(
        select(StudyTemplate).where(StudyTemplate.id == template_id)
    )
    template = template_result.scalar_one_or_none()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Check if template has DOCX file
    if not hasattr(template, 'template_file_path') or not template.template_file_path:
        raise HTTPException(
            status_code=400,
            detail="Template does not have a DOCX file."
        )
    
    # Get user info
    user_id = current_user.get("user_id", "anonymous") if current_user else "anonymous"
    user_name = current_user.get("name", "User") if current_user else "User"
    
    # Create document URL using configured backend internal URL
    internal_base = settings.backend_internal_url
    document_url = f"{internal_base}/api/templates/{template_id}/template-file"
    
    # Generate document key (unique identifier for ONLYOFFICE)
    document_key = str(template.id)
    
    # Create configuration in view mode (read-only for templates)
    config = create_document_config(
        document_url=document_url,
        callback_url="",  # No callback needed for template viewing
        document_key=document_key,
        document_title=template.template_name,
        user_id=user_id,
        user_name=user_name,
        mode="view"  # Always view mode for templates
    )
    
    logger.info(f"ONLYOFFICE template config - document_url: {document_url}, template_id: {template_id}")
    
    return JSONResponse(content={
        "editorUrl": get_onlyoffice_editor_url(),
        "config": config
    })


@router.patch("/templates/{template_id}/deactivate")
async def deactivate_template(
    template_id: UUID,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Deactivate a template (soft delete).
    Deactivated templates cannot be used for new agreements.
    """
    template_result = await db.execute(
        select(StudyTemplate).where(StudyTemplate.id == template_id)
    )
    template = template_result.scalar_one_or_none()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    template.is_active = 'false'
    await db.flush()
    await db.commit()
    
    return {"message": "Template deactivated successfully"}


@router.put("/templates/{template_id}/placeholder-config", response_model=schemas.StudyTemplateResponse)
async def update_template_placeholder_config(
    template_id: UUID,
    config_data: schemas.PlaceholderConfigUpdate,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Update placeholder configuration for a template.
    Allows setting which placeholders are editable in agreements created from this template.
    """
    template_result = await db.execute(
        select(StudyTemplate).where(StudyTemplate.id == template_id)
    )
    template = template_result.scalar_one_or_none()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Update placeholder configuration
    template.placeholder_config = config_data.placeholder_config
    await db.flush()
    await db.commit()
    await db.refresh(template)
    
    return {
        "id": template.id,
        "study_id": template.study_id,
        "template_name": template.template_name,
        "template_type": template.template_type.value,
        "template_content": template.template_content if hasattr(template, 'template_content') and template.template_content else {"type": "doc", "content": []},
        "template_file_path": template.template_file_path if hasattr(template, 'template_file_path') else None,
        "placeholder_config": template.placeholder_config if hasattr(template, 'placeholder_config') and template.placeholder_config else {},
        "field_mappings": template.field_mappings if hasattr(template, 'field_mappings') and template.field_mappings else None,
        "created_by": template.created_by,
        "created_at": template.created_at,
        "updated_at": template.updated_at,
        "is_active": template.is_active,
    }


@router.put("/templates/{template_id}/field-mappings", response_model=schemas.StudyTemplateResponse)
async def update_template_field_mappings(
    template_id: UUID,
    mappings_data: schemas.FieldMappingsUpdate,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    db: AsyncSession = Depends(get_db),
):
    """
    Update field mappings for a template.
    Allows configuring which data sources map to template placeholders.
    Format: {"PLACEHOLDER_NAME": "data_source.field_name"}
    Examples:
    - {"SITE_NAME": "site_profile.site_name"}
    - {"PI_NAME": "site_profile.pi_name"}
    - {"STUDY_TITLE": "agreement.title"}
    """
    template_result = await db.execute(
        select(StudyTemplate).where(StudyTemplate.id == template_id)
    )
    template = template_result.scalar_one_or_none()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Validate field mappings format
    for placeholder_name, mapping_path in mappings_data.field_mappings.items():
        if not isinstance(placeholder_name, str) or not isinstance(mapping_path, str):
            raise HTTPException(
                status_code=400,
                detail=f"Invalid field_mappings format. All keys and values must be strings."
            )
        # Validate mapping path format (should be "data_source.field_name")
        parts = mapping_path.split(".", 1)
        if len(parts) != 2:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid mapping path format for {placeholder_name}: {mapping_path}. Expected format: 'data_source.field_name'"
            )
        data_source = parts[0].lower()
        if data_source not in ["site_profile", "agreement"]:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid data source for {placeholder_name}: {data_source}. Must be 'site_profile' or 'agreement'"
            )
    
    # Update field mappings
    template.field_mappings = mappings_data.field_mappings
    logger.info(f"Updated field_mappings for template {template_id}: {mappings_data.field_mappings}")
    await db.flush()
    await db.commit()
    await db.refresh(template)
    
    return {
        "id": template.id,
        "study_id": template.study_id,
        "template_name": template.template_name,
        "template_type": template.template_type.value,
        "template_content": template.template_content if hasattr(template, 'template_content') and template.template_content else {"type": "doc", "content": []},
        "template_file_path": template.template_file_path if hasattr(template, 'template_file_path') else None,
        "placeholder_config": template.placeholder_config if hasattr(template, 'placeholder_config') and template.placeholder_config else {},
        "field_mappings": template.field_mappings if hasattr(template, 'field_mappings') and template.field_mappings else None,
        "created_by": template.created_by,
        "created_at": template.created_at,
        "updated_at": template.updated_at,
        "is_active": template.is_active,
    }
