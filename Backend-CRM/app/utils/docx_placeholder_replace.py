"""
DOCX Placeholder Replacement Module

This module handles placeholder replacement in DOCX documents with support for:
- Dynamic field mappings
- Editable and locked placeholders
- Content control creation for editable fields
- Document protection
"""

import logging
import copy
from pathlib import Path
from typing import Optional, Dict, Any


from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement

logger = logging.getLogger(__name__)


def get_placeholder_value_from_mapping(
    placeholder_name: str,
    field_mappings: Dict[str, str],
    site_profile=None,
    agreement=None,
) -> Optional[str]:
    """
    Get the value for a placeholder using dynamic field mappings.
    
    Args:
        placeholder_name: The placeholder name (e.g., "SITE_NAME")
        field_mappings: Dict mapping placeholder names to data source paths (e.g., {"SITE_NAME": "site_profile.site_name"})
        site_profile: The SiteProfile instance
        agreement: The Agreement instance
        
    Returns:
        The value from the data source as a string, or None if:
        - No explicit field mapping exists for this placeholder
        - The data source object is missing
        - The resolved field value is null
        
        IMPORTANT: Placeholders without explicit field mappings will remain unchanged.
        No fallback to legacy mapping is used.
    """
    mapping_path: Optional[str] = None

    # Only use explicit field_mappings from template configuration
    # NO FALLBACK to legacy mapping - placeholders without mapping should remain unchanged
    if field_mappings:
        mapping_path = field_mappings.get(placeholder_name.upper())
        if not mapping_path:
            logger.debug(
                "No explicit field_mappings entry for placeholder '%s'; "
                "placeholder will remain unchanged.",
                placeholder_name,
            )
            return None
    else:
        logger.debug(
            "No field_mappings provided for placeholder '%s'; "
            "placeholder will remain unchanged.",
            placeholder_name,
        )
        return None
    
    # Parse the mapping path (format: "data_source.field_name")
    parts = mapping_path.split(".", 1)
    if len(parts) != 2:
        logger.warning(
            "Invalid mapping path format for placeholder %s: %s. Expected format: 'data_source.field_name'",
            placeholder_name,
            mapping_path,
        )
        return None
    
    data_source, field_name = parts[0].lower(), parts[1]
    
    # Get value from appropriate data source
    value = None
    if data_source == "site_profile":
        if not site_profile:
            logger.warning("SiteProfile not provided for placeholder: %s", placeholder_name)
            return None
        value = getattr(site_profile, field_name, None)
    elif data_source == "agreement":
        if not agreement:
            logger.warning("Agreement not provided for placeholder: %s", placeholder_name)
            return None
        value = getattr(agreement, field_name, None)
    else:
        logger.warning(
            "Unsupported data source '%s' for placeholder '%s'. Supported: site_profile, agreement",
            data_source,
            placeholder_name,
        )
        return None
    
    if value is None:
        logger.debug(
            "Field '%s' is null in %s for placeholder '%s'",
            field_name,
            data_source,
            placeholder_name,
        )
        return None
    
    # Convert to string for text replacement
    return str(value) if value is not None else ""


def replace_placeholders_in_docx(
    docx_path: Path,
    profile,
    output_path: Path,
    sponsor_signatory_name: str,
    sponsor_signatory_email: str,
    current_user_email: str,
    field_mappings: Optional[Dict[str, str]] = None,
    agreement=None,
    template_id: Optional[str] = None,
    placeholder_config: Optional[Dict[str, Dict[str, Any]]] = None,
) -> Path:
    """
    Replace placeholders in a DOCX document using dynamic field mappings.
    
    Args:
        docx_path: Path to the input DOCX file
        profile: SiteProfile object containing data
        output_path: Path where the modified DOCX will be saved
        sponsor_signatory_name: Name of the sponsor signatory
        sponsor_signatory_email: Email of the sponsor signatory
        current_user_email: Email of the current user
        field_mappings: Dict mapping placeholder names to data source paths
        agreement: Optional Agreement object
        template_id: Optional template ID for logging
        placeholder_config: Optional configuration for placeholder locking/editing
        
    Returns:
        Path to the output file
        
    Raises:
        Exception: If document processing fails
    """
    try:
        # Load the document
        doc = Document(docx_path)
        
        # Determine which mapping to use
        use_dynamic_mapping = field_mappings is not None and len(field_mappings) > 0
        
        if use_dynamic_mapping:
            logger.info(f"Using dynamic field_mappings: {field_mappings}")
            # Extract all placeholders from the document
            placeholders_found = _extract_placeholders_from_docx(doc)
            logger.info(f"Found {len(placeholders_found)} unique placeholders in document")
            
            # Create a mapping of placeholder text to replacement value
            placeholder_replacements = {}
            for placeholder_name in placeholders_found:
                # Skip signature blocks - handled separately
                if placeholder_name in ["SITE_SIGNATURE_BLOCK", "SPONSOR_SIGNATURE_BLOCK"]:
                    continue
                
                value = get_placeholder_value_from_mapping(
                    placeholder_name,
                    field_mappings,
                    site_profile=profile,
                    agreement=agreement,
                )
                
                placeholder_text = f"{{{{{placeholder_name}}}}}"
                
                # Check if this placeholder is editable
                is_editable = True  # Default to editable if no config
                if placeholder_config:
                    placeholder_cfg = placeholder_config.get(placeholder_name, {})
                    is_editable = placeholder_cfg.get("editable", True)
                
                if value is not None:
                    placeholder_replacements[placeholder_text] = {
                        'value': value,
                        'editable': is_editable,
                        'name': placeholder_name
                    }
                    logger.debug(f"Found replacement for {placeholder_name}: {value}")
                else:
                    logger.debug(f"No value found for placeholder {placeholder_name}")
            
            # Apply replacements in paragraphs
            replacements_made = 0
            for paragraph in doc.paragraphs:
                replacements_made += _replace_placeholders_in_paragraph(
                    paragraph, placeholder_replacements
                )
            
            # Apply replacements in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replacements_made += _replace_placeholders_in_paragraph(
                                paragraph, placeholder_replacements
                            )
            
            # Apply document protection for locked fields
            if placeholder_config:
                _apply_restricted_editing_protection(doc)
            
            logger.info(f"Completed placeholder replacement with {replacements_made} replacements")
        
        else:
            # Legacy fallback - simple text replacement
            logger.warning("No field_mappings provided, using legacy placeholder replacement")
            _replace_legacy_placeholders(doc, profile)
        
        # Handle signature blocks (always use legacy logic)
        _replace_signature_blocks_in_docx(
            doc,
            profile,
            sponsor_signatory_name,
            sponsor_signatory_email,
            current_user_email
        )
        
        # Apply document protection for templates that have placeholder_config
        if placeholder_config is not None:
            editable_field_names = [
                name for name, cfg in placeholder_config.items()
                if cfg.get("editable", True)
            ]
            logger.info(
                "Applying document protection for template_id=%s. "
                "Editable placeholders: %s",
                template_id,
                ", ".join(editable_field_names) if editable_field_names else "none",
            )
            try:
                _apply_restricted_editing_protection(doc)
                logger.info("Document protection applied successfully")
            except Exception as e:
                logger.error(
                    "Failed to apply document protection; document will remain fully editable. "
                    "Error: %s",
                    e,
                    exc_info=True,
                )
        
        # Save the modified document
        doc.save(output_path)
        logger.info(f"Successfully replaced placeholders in DOCX: {output_path}")
        
        return output_path
        
    except Exception as e:
        logger.error(f"Failed to replace placeholders in DOCX: {str(e)}", exc_info=True)
        raise Exception(f"Failed to process DOCX file: {str(e)}")


def _extract_placeholders_from_docx(doc: Document) -> set:
    """
    Extract all placeholder names from a DOCX document.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        Set of placeholder names (without {{ }})
    """
    import re
    placeholders = set()
    pattern = r'\{\{([A-Z0-9_]+)\}\}'
    
    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        placeholders.update(matches)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    placeholders.update(matches)
    
    return placeholders


def _apply_restricted_editing_protection(doc: Document) -> None:
    """
    Apply Word document protection to restrict editing to content controls only.
    
    This makes the entire document read-only except for editable content controls.
    Users can only edit text inside content controls; all other text is protected.
    
    Args:
        doc: Document object to protect
    """
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import OxmlElement
        
        # Get settings part from related parts
        settings_part = None
        for rel in doc.part.rels.values():
            if rel.reltype == RT.SETTINGS:
                settings_part = rel.target_part
                break
        
        if settings_part is None:
            logger.warning("Settings part not found - document protection cannot be applied. Document will work but editing may not be restricted.")
            return
        
        settings_xml = settings_part.element
        
        # Check if documentProtection already exists
        doc_protection = settings_xml.find(qn('w:documentProtection'))
        
        if doc_protection is None:
            # Create new documentProtection element
            doc_protection = OxmlElement('w:documentProtection')
            settings_xml.append(doc_protection)
        
        # Set protection attributes
        doc_protection.set(qn('w:enforcement'), '1')
        doc_protection.set(qn('w:edit'), 'forms')  # "forms" allows editing only in content controls
        doc_protection.set(qn('w:cryptProviderType'), 'rsaFull')
        doc_protection.set(qn('w:cryptAlgorithmClass'), 'hash')
        doc_protection.set(qn('w:cryptAlgorithmType'), 'typeAny')
        doc_protection.set(qn('w:cryptAlgorithmSid'), '4')
        doc_protection.set(qn('w:cryptSpinCount'), '100000')
        doc_protection.set(qn('w:hash'), '')
        
        logger.info("Applied document protection: restricted editing mode (editing allowed only in content controls)")
        
    except Exception as e:
        logger.error(f"Failed to apply document protection: {str(e)}", exc_info=True)
        # Don't fail the entire operation if protection fails


def _replace_text_across_runs(paragraph, old_text: str, new_text: str) -> bool:
    """
    Replace text in a paragraph that may span multiple XML runs.
    
    Word often splits what looks like continuous text (e.g., '{{CITY}}') across
    multiple runs due to formatting, spell-check, or editing history. This function
    handles that by working on the concatenated text of all runs, then distributing
    the result back.
    
    Args:
        paragraph: The paragraph object
        old_text: The text to find and replace
        new_text: The replacement text
        
    Returns:
        True if replacement was made, False otherwise
    """
    runs = paragraph.runs
    if not runs:
        return False
    
    # Build the full text from all runs
    full_text = ''.join(run.text for run in runs)
    
    start_idx = full_text.find(old_text)
    if start_idx == -1:
        return False
    
    end_idx = start_idx + len(old_text)
    
    # Build a char-position-to-run mapping
    run_boundaries = []  # list of (run_index, start_char, end_char)
    char_pos = 0
    for i, run in enumerate(runs):
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_boundaries.append((i, run_start, run_end))
        char_pos = run_end
    
    # Find which runs the placeholder starts and ends in
    start_run_idx = None
    end_run_idx = None
    for (idx, rs, re_) in run_boundaries:
        if start_run_idx is None and start_idx < re_:
            start_run_idx = idx
        if end_idx <= re_:
            end_run_idx = idx
            break
    
    if start_run_idx is None or end_run_idx is None:
        return False
    
    # Calculate offsets within the start and end runs
    start_run_char_start = run_boundaries[start_run_idx][1]
    end_run_char_start = run_boundaries[end_run_idx][1]
    
    offset_in_start_run = start_idx - start_run_char_start
    offset_in_end_run = end_idx - end_run_char_start
    
    if start_run_idx == end_run_idx:
        # Simple case: placeholder is within a single run
        run = runs[start_run_idx]
        run.text = run.text[:offset_in_start_run] + new_text + run.text[offset_in_end_run:]
    else:
        # Placeholder spans multiple runs
        # Put everything before placeholder + replacement in start run
        runs[start_run_idx].text = runs[start_run_idx].text[:offset_in_start_run] + new_text
        
        # Put everything after placeholder in end run
        runs[end_run_idx].text = runs[end_run_idx].text[offset_in_end_run:]
        
        # Clear all intermediate runs
        for i in range(start_run_idx + 1, end_run_idx):
            runs[i].text = ''
    
    return True


def _replace_placeholders_in_paragraph(paragraph, placeholder_replacements: dict) -> int:
    """
    Replace all placeholders in a paragraph with their corresponding values.
    Handles placeholders split across multiple XML runs.
    
    Args:
        paragraph: The paragraph object to modify
        placeholder_replacements: Dict mapping placeholder text to replacement info
        
    Returns:
        Number of replacements made
    """
    replacements_made = 0
    
    # Get all placeholder texts that exist in this paragraph's full text
    full_text = paragraph.text
    found_placeholders = [pt for pt in placeholder_replacements.keys() if pt in full_text]
    
    if not found_placeholders:
        return 0
    
    # Sort placeholders by length (longest first) to avoid partial replacements
    found_placeholders.sort(key=len, reverse=True)
    
    # Separate editable and locked placeholders
    editable_placeholders = []
    locked_placeholders = []
    
    for placeholder_text in found_placeholders:
        if placeholder_text in placeholder_replacements:
            replacement_info = placeholder_replacements[placeholder_text]
            value = replacement_info['value']
            is_editable = replacement_info['editable']
            placeholder_name = replacement_info['name']
            
            if is_editable:
                # For editable fields, track for content control creation
                logger.debug(f"Marking editable placeholder {placeholder_name} with: {value}")
                editable_placeholders.append({
                    'text': placeholder_text,
                    'value': value,
                    'name': placeholder_name
                })
            else:
                # For locked fields, replace with plain text
                logger.debug(f"Replacing locked placeholder {placeholder_name} with: {value}")
                locked_placeholders.append({
                    'text': placeholder_text,
                    'value': value,
                    'name': placeholder_name
                })
    
    # Replace locked placeholders with plain text
    for placeholder_info in locked_placeholders:
        placeholder_text = placeholder_info['text']
        value = placeholder_info['value']
        replaced = _replace_text_across_runs(paragraph, placeholder_text, value)
        if replaced:
            replacements_made += 1
        else:
            logger.warning(f"Failed to replace locked placeholder {placeholder_info['name']} in paragraph")
    
    # Create editable content controls for editable placeholders
    if editable_placeholders:
        _create_editable_content_controls_in_paragraph(paragraph, editable_placeholders)
        replacements_made += len(editable_placeholders)
    
    return replacements_made


# def _create_editable_content_controls_in_paragraph(paragraph, editable_placeholders):
#     """
#     Create actual editable content controls in a paragraph using Word XML.
    
#     Args:
#         paragraph: The paragraph object
#         editable_placeholders: List of dicts with 'text', 'value', and 'name' keys
#     """
#     try:
#         # For editable placeholders, use the same cross-run replacement as locked fields
#         # The document protection will handle making them editable
#         for placeholder_info in editable_placeholders:
#             placeholder_text = placeholder_info['text']
#             value = placeholder_info['value']
#             placeholder_name = placeholder_info['name']
            
#             # Use cross-run replacement to preserve formatting
#             replaced = _replace_text_across_runs(paragraph, placeholder_text, value)
#             if replaced:
#                 logger.debug(f"Replaced editable placeholder {placeholder_name} with: {value}")
#             else:
#                 logger.warning(f"Failed to replace editable placeholder {placeholder_name} in paragraph")
                        
#     except Exception as e:
#         logger.error(f"Error creating editable content control: {e}")
#         # Fallback: just replace the text
#         for placeholder_info in editable_placeholders:
#             placeholder_text = placeholder_info['text']
#             value = placeholder_info['value']
#             paragraph.text = paragraph.text.replace(placeholder_text, value)

def _create_editable_content_controls_in_paragraph(paragraph, editable_placeholders):
    """
    Create actual editable content controls in a paragraph using Word XML.
    This allows the text to remain editable even when document protection is applied,
    while strictly preserving the original formatting.
    """
    try:
        for placeholder_info in editable_placeholders:
            placeholder_text = placeholder_info['text']
            value = placeholder_info['value']
            placeholder_name = placeholder_info['name']
            
            # Step 1: Replace placeholder with a unique token so it lands perfectly inside ONE run
            unique_token = f"@@@_{placeholder_name}_@@@"
            replaced = _replace_text_across_runs(paragraph, placeholder_text, unique_token)
            
            if not replaced:
                continue
                
            # Step 2: Locate the specific run containing our unique token
            for run in paragraph.runs:
                if unique_token in run.text:
                    # Split the text around the token to maintain text before/after
                    before_text, after_text = run.text.split(unique_token, 1)
                    
                    # Keep the 'before' text in the original run
                    run.text = before_text
                    
                    # Step 3: Create the Structured Document Tag (Content Control)
                    sdt = OxmlElement('w:sdt')
                    
                    # Content Control Properties
                    sdtPr = OxmlElement('w:sdtPr')
                    alias = OxmlElement('w:alias')
                    alias.set(qn('w:val'), placeholder_name)
                    tag = OxmlElement('w:tag')
                    tag.set(qn('w:val'), placeholder_name)
                    
                    sdtPr.append(alias)
                    sdtPr.append(tag)
                    
                    # Specify that this is a plain text content control
                    text_pr = OxmlElement('w:text')
                    sdtPr.append(text_pr)
                    sdt.append(sdtPr)
                    
                    # Content Control Content
                    sdtContent = OxmlElement('w:sdtContent')
                    sdt.append(sdtContent)
                    
                    # Step 4: Create a new run for the placeholder value to sit inside the SDT
                    new_run_sdt = OxmlElement('w:r')
                    
                    # CRITICAL: Deep copy formatting (rPr) from the original run so styling perfectly aligns
                    if run._r.rPr is not None:
                        new_run_sdt.append(copy.deepcopy(run._r.rPr))
                        
                    new_run_text = OxmlElement('w:t')
                    new_run_text.text = value
                    if value.startswith(' ') or value.endswith(' '):
                        new_run_text.set(qn('xml:space'), 'preserve')
                    new_run_sdt.append(new_run_text)
                    
                    sdtContent.append(new_run_sdt)
                    
                    # Insert the SDT node immediately after the original run
                    run._r.addnext(sdt)
                    
                    # Step 5: If there was text after the token, create another run for it to prevent layout breaks
                    if after_text:
                        new_run_after = OxmlElement('w:r')
                        if run._r.rPr is not None:
                            new_run_after.append(copy.deepcopy(run._r.rPr))
                        new_run_after_text = OxmlElement('w:t')
                        new_run_after_text.text = after_text
                        if after_text.startswith(' ') or after_text.endswith(' '):
                            new_run_after_text.set(qn('xml:space'), 'preserve')
                        new_run_after.append(new_run_after_text)
                        
                        sdt.addnext(new_run_after)
                        
                    logger.debug(f"Created true editable content control for {placeholder_name}")
                    break
                    
    except Exception as e:
        logger.error(f"Error creating editable content control: {e}")
        # Fallback: Just drop the plain text in so the document isn't ruined
        for placeholder_info in editable_placeholders:
            placeholder_text = placeholder_info['text']
            value = placeholder_info['value']
            # Fallback will lose the editable property under protection, but text will render
            _replace_text_across_runs(paragraph, placeholder_text, value)


def _replace_legacy_placeholders(doc: Document, profile) -> None:
    """Legacy placeholder replacement for backward compatibility.
    Handles placeholders split across multiple DOCX XML runs."""
    legacy_mappings = {
        'SITE_NAME': getattr(profile, 'site_name', ''),
        'HOSPITAL_NAME': getattr(profile, 'hospital_name', ''),
        'CITY': getattr(profile, 'city', ''),
        'STATE': getattr(profile, 'state', ''),
        'COUNTRY': getattr(profile, 'country', ''),
        'POSTAL_CODE': getattr(profile, 'postal_code', ''),
        'ADDRESS_LINE_1': getattr(profile, 'address_line_1', ''),
        'ADDRESS_LINE_2': getattr(profile, 'address_line_2', ''),
        'PI_NAME': getattr(profile, 'pi_name', ''),
        'PI_EMAIL': getattr(profile, 'pi_email', ''),
        'PI_PHONE': getattr(profile, 'pi_phone', ''),
        'FULL_ADDRESS': getattr(profile, 'full_address', ''),
    }
    
    def _replace_in_paragraph(paragraph):
        for placeholder, value in legacy_mappings.items():
            if value:
                placeholder_text = f"{{{{{placeholder}}}}}"
                if placeholder_text in paragraph.text:
                    _replace_text_across_runs(paragraph, placeholder_text, str(value))
    
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph)


def _replace_signature_blocks_in_docx(
    doc: Document,
    profile,
    sponsor_signatory_name: str,
    sponsor_signatory_email: str,
    current_user_email: str,
) -> None:
    """Replace signature block placeholders"""
    # Handle signature blocks
    for paragraph in doc.paragraphs:
        if "{{SITE_SIGNATURE_BLOCK}}" in paragraph.text:
            # Create site signature block
            signature_text = f"""
            _______________________________
            {getattr(profile, 'pi_name', 'Principal Investigator')}
            {getattr(profile, 'pi_email', '')}
            Date: ___________
            """
            _replace_text_across_runs(paragraph, "{{SITE_SIGNATURE_BLOCK}}", signature_text.strip())
        
        if "{{SPONSOR_SIGNATURE_BLOCK}}" in paragraph.text:
            # Create sponsor signature block
            signature_text = f"""
            _______________________________
            {sponsor_signatory_name}
            {sponsor_signatory_email}
            Date: ___________
            """
            _replace_text_across_runs(paragraph, "{{SPONSOR_SIGNATURE_BLOCK}}", signature_text.strip())

